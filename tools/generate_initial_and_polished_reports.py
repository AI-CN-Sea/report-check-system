from pathlib import Path
import math
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSET_DIR = DOCS / "report_assets_v2"
ASSET_DIR.mkdir(exist_ok=True)

TITLE = "基于 Spring Boot 与文本相似度算法的实验报告查重系统设计与实现"
BASE_NAME = "学号_姓名_基于Spring Boot与文本相似度算法的实验报告查重系统设计与实现"
INITIAL_OUT = DOCS / f"{BASE_NAME}_论文初稿.docx"
FINAL_OUT = DOCS / f"{BASE_NAME}_论文终稿.docx"


def font_path(name):
    p = Path("C:/Windows/Fonts") / name
    return str(p) if p.exists() else None


CN_FONT = font_path("simsun.ttc") or font_path("msyh.ttc") or font_path("arial.ttf")


def img_font(size):
    try:
        return ImageFont.truetype(CN_FONT, size)
    except Exception:
        return ImageFont.load_default()


INK = (35, 48, 64)
BLUE = (44, 88, 140)
MUTED = (102, 116, 132)
LINE = (92, 119, 150)


def text_box(draw, text, box, font, fill=INK, align="center", spacing=6):
    x1, y1, x2, y2 = box
    max_w = x2 - x1 - 22
    lines = []
    for para in text.split("\n"):
        line = ""
        for ch in para:
            trial = line + ch
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_w:
                line = trial
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
    heights = [draw.textbbox((0, 0), line, font=font)[3] - draw.textbbox((0, 0), line, font=font)[1] for line in lines]
    total_h = sum(heights) + max(0, len(lines) - 1) * spacing
    y = y1 + max(0, (y2 - y1 - total_h) / 2)
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font)[2]
        x = x1 + 14 if align == "left" else x1 + (x2 - x1 - w) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + spacing


def card(draw, box, text, fill, font, outline=(155, 171, 190), radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=2)
    text_box(draw, text, box, font)


def arrow(draw, start, end, fill=LINE, width=4):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    ang = math.atan2(ey - sy, ex - sx)
    size = 14
    pts = [
        (ex, ey),
        (ex - size * math.cos(ang - math.pi / 6), ey - size * math.sin(ang - math.pi / 6)),
        (ex - size * math.cos(ang + math.pi / 6), ey - size * math.sin(ang + math.pi / 6)),
    ]
    draw.polygon(pts, fill=fill)


def title(draw, text, width):
    f = img_font(42)
    w = draw.textbbox((0, 0), text, font=f)[2]
    draw.text(((width - w) / 2, 34), text, font=f, fill=BLUE)


def make_architecture(path):
    img = Image.new("RGB", (1700, 980), "white")
    d = ImageDraw.Draw(img)
    title(d, "系统总体架构", 1700)
    f, sf = img_font(27), img_font(22)
    # layer backgrounds
    bands = [
        (80, 140, 1620, 300, "#F4F8FC", "表现层"),
        (80, 355, 1620, 555, "#F7FAF5", "服务层"),
        (80, 610, 1620, 820, "#FBF7F2", "数据与算法支撑层"),
    ]
    for x1, y1, x2, y2, fill, label in bands:
        d.rounded_rectangle((x1, y1, x2, y2), radius=26, fill=fill, outline=(218, 226, 235), width=2)
        d.text((x1 + 22, y1 + 18), label, font=sf, fill=MUTED)
    boxes = [
        ((250, 180, 520, 265), "浏览器访问\n教师/学生/管理员", "#FFFFFF"),
        ((700, 180, 1000, 265), "Vue 3 前端\nElement Plus / ECharts", "#FFFFFF"),
        ((350, 410, 650, 510), "Spring Boot API\nController / Service", "#FFFFFF"),
        ((810, 410, 1110, 510), "权限与业务处理\nJWT / 角色控制", "#FFFFFF"),
        ((1260, 410, 1530, 510), "文件处理\n上传 / 解析 / 预览", "#FFFFFF"),
        ((220, 670, 500, 760), "MySQL 数据库\n用户/任务/结果", "#FFFFFF"),
        ((700, 670, 1010, 760), "本地文件存储\nuploads 目录", "#FFFFFF"),
        ((1200, 670, 1510, 760), "查重算法模块\nTF-IDF / SimHash", "#FFFFFF"),
    ]
    for box, txt, fill in boxes:
        card(d, box, txt, fill, f)
    arrow(d, (520, 222), (700, 222))
    arrow(d, (850, 265), (500, 410))
    arrow(d, (850, 265), (960, 410))
    arrow(d, (1000, 222), (1395, 410))
    arrow(d, (500, 510), (360, 670))
    arrow(d, (1395, 510), (855, 670))
    arrow(d, (960, 510), (1355, 670))
    img.save(path)


def make_function(path):
    img = Image.new("RGB", (1700, 1000), "white")
    d = ImageDraw.Draw(img)
    title(d, "系统功能结构", 1700)
    f, sf = img_font(25), img_font(21)
    card(d, (610, 115, 1090, 190), "实验报告智能查重与分析系统", "#EAF2FB", f, outline=(120, 150, 185))
    columns = [
        (120, "管理员端", ["用户管理", "课程管理", "班级管理", "班级学生分配", "系统看板"], "#F6F8FB"),
        (520, "教师端", ["实验任务管理", "报告上传与批量导入", "发起查重任务", "查看查重结果", "统计分析与导出"], "#F5FAF4"),
        (920, "学生端", ["查看实验任务", "提交实验报告", "查看提交状态"], "#FAF7F2"),
        (1220, "算法与数据", ["文本解析", "文本预处理", "TF-IDF 余弦相似度", "SimHash 指纹", "风险等级与相似句"], "#F8F4FA"),
    ]
    for x, head, items, fill in columns:
        card(d, (x, 280, x + 280, 350), head, fill, f, outline=(145, 160, 180))
        arrow(d, (850, 190), (x + 140, 280), width=3)
        y = 400
        for item in items:
            card(d, (x, y, x + 280, y + 55), item, "#FFFFFF", sf, outline=(190, 199, 210), radius=12)
            y += 74
    img.save(path)


def make_algorithm(path):
    img = Image.new("RGB", (1750, 950), "white")
    d = ImageDraw.Draw(img)
    title(d, "查重算法流程", 1750)
    f, sf = img_font(23), img_font(25)
    steps = [
        ("报告集合", "读取同一实验任务下\n全部已解析报告"),
        ("文本预处理", "清洗文本、二元词条切分\n停用词过滤"),
        ("权重建模", "统计文档频率\n计算 TF-IDF 向量"),
        ("相似度计算", "两两计算余弦相似度\n生成 SimHash 指纹"),
        ("融合判定", "0.7×余弦 + 0.3×SimHash\n划分风险等级"),
        ("结果解释", "保存查重结果\n提取相似句详情"),
    ]
    x0, y0, w, h, gap = 100, 250, 230, 130, 42
    for i, (head, body) in enumerate(steps):
        x = x0 + i * (w + gap)
        card(d, (x, y0, x + w, y0 + h), f"{head}\n{body}", "#F3F8FC" if i % 2 == 0 else "#FBF6EF", f)
        if i < len(steps) - 1:
            arrow(d, (x + w, y0 + h / 2), (x + w + gap, y0 + h / 2))
    d.rounded_rectangle((280, 560, 1470, 740), radius=24, fill="#F6F7F9", outline=(160, 174, 190), width=2)
    text_box(d, "输出结果并非直接判定抄袭，而是为教师提供综合相似度、风险等级和相似句等复核依据。", (340, 595, 1410, 700), sf)
    img.save(path)


def make_er(path):
    img = Image.new("RGB", (1750, 1080), "white")
    d = ImageDraw.Draw(img)
    title(d, "数据库核心关系", 1750)
    f, sf = img_font(22), img_font(18)
    nodes = {
        "sys_role\n角色": (120, 170, 350, 270),
        "sys_user\n用户": (500, 170, 730, 270),
        "class_info\n班级": (880, 170, 1110, 270),
        "student_class\n学生班级关联": (1260, 170, 1530, 270),
        "course\n课程": (500, 420, 730, 520),
        "experiment_task\n实验任务": (880, 420, 1150, 520),
        "report_file\n报告文件": (1260, 420, 1530, 520),
        "check_task\n查重任务": (880, 690, 1150, 790),
        "check_result\n查重结果": (1260, 690, 1530, 790),
        "similar_sentence\n相似句子": (1260, 900, 1530, 1000),
    }
    for text, box in nodes.items():
        card(d, box, text, "#F7FAFD", f, outline=(130, 155, 185), radius=14)
    def c(box):
        return ((box[0] + box[2]) / 2, (box[1] + box[3]) / 2)
    for a, b in [
        ("sys_role\n角色", "sys_user\n用户"),
        ("sys_user\n用户", "student_class\n学生班级关联"),
        ("class_info\n班级", "student_class\n学生班级关联"),
        ("sys_user\n用户", "course\n课程"),
        ("course\n课程", "experiment_task\n实验任务"),
        ("class_info\n班级", "experiment_task\n实验任务"),
        ("experiment_task\n实验任务", "report_file\n报告文件"),
        ("experiment_task\n实验任务", "check_task\n查重任务"),
        ("report_file\n报告文件", "check_result\n查重结果"),
        ("check_task\n查重任务", "check_result\n查重结果"),
        ("check_result\n查重结果", "similar_sentence\n相似句子"),
    ]:
        arrow(d, c(nodes[a]), c(nodes[b]), width=3)
    d.text((110, 1015), "说明：operation_log 表作为扩展表预留，可继续记录登录、上传、查重等关键行为。", font=sf, fill=MUTED)
    img.save(path)


def make_tech_flow(path):
    img = Image.new("RGB", (1750, 960), "white")
    d = ImageDraw.Draw(img)
    title(d, "系统技术流程", 1750)
    f = img_font(24)
    lanes = [
        (120, 170, 1620, 300, "#F4F8FC", "前端交互"),
        (120, 350, 1620, 520, "#F6FAF4", "后端业务"),
        (120, 570, 1620, 760, "#FBF7F2", "数据与算法"),
    ]
    for x1, y1, x2, y2, fill, label in lanes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline=(220, 226, 235), width=2)
        d.text((x1 + 22, y1 + 18), label, font=img_font(22), fill=MUTED)
    items = [
        ((280, 215, 530, 280), "选择任务/上传报告"),
        ((760, 215, 1010, 280), "查看结果/统计图表"),
        ((280, 405, 530, 480), "接口鉴权与参数校验"),
        ((760, 405, 1010, 480), "文件保存与文本解析"),
        ((1240, 405, 1490, 480), "创建查重任务"),
        ((280, 635, 530, 710), "MySQL 保存业务数据"),
        ((760, 635, 1010, 710), "TF-IDF/SimHash 计算"),
        ((1240, 635, 1490, 710), "写入结果与相似句"),
    ]
    for box, txt in items:
        card(d, box, txt, "#FFFFFF", f, outline=(170, 188, 205), radius=14)
    for s, e in [((530, 248), (760, 248)), ((405, 280), (405, 405)), ((885, 280), (885, 405)), ((530, 442), (760, 442)), ((1010, 442), (1240, 442)), ((405, 480), (405, 635)), ((885, 480), (885, 635)), ((1365, 480), (1365, 635)), ((530, 672), (760, 672)), ((1010, 672), (1240, 672)), ((1365, 635), (885, 300))]:
        arrow(d, s, e, width=3)
    img.save(path)


def make_progress(path):
    img = Image.new("RGB", (1700, 920), "white")
    d = ImageDraw.Draw(img)
    title(d, "实验进度与成果", 1700)
    f, sf = img_font(24), img_font(20)
    stages = [
        ("阶段一\n需求与开题", "完成开题报告\n确定技术路线"),
        ("阶段二\n基础框架", "后端/前端/数据库\n完成基础搭建"),
        ("阶段三\n核心功能", "报告上传解析\n查重算法实现"),
        ("阶段四\n演示数据", "多班级多任务样例\n批量上传验证"),
        ("阶段五\n测试与论文", "单元测试/构建验证\n整理报告终稿"),
    ]
    x, y, w, h, gap = 115, 330, 250, 140, 70
    d.line((200, y + h + 65, 1500, y + h + 65), fill=(180, 190, 205), width=6)
    for i, (head, body) in enumerate(stages):
        bx = x + i * (w + gap)
        card(d, (bx, y, bx + w, y + h), f"{head}\n{body}", "#F7FAFD" if i % 2 == 0 else "#FAF7F2", f, radius=18)
        cx = bx + w / 2
        d.ellipse((cx - 18, y + h + 47, cx + 18, y + h + 83), fill="#2C588C")
        d.text((cx - 8, y + h + 51), str(i + 1), font=sf, fill="white")
    text_box(d, "当前系统已完成可运行闭环：登录 → 任务 → 上传 → 解析 → 查重 → 结果 → 图表 → 导出。", (260, 650, 1440, 760), img_font(28))
    img.save(path)


def make_business(path):
    img = Image.new("RGB", (1700, 940), "white")
    d = ImageDraw.Draw(img)
    title(d, "教师端查重业务流程", 1700)
    f, sf = img_font(24), img_font(21)
    lanes = [
        (95, 155, 1605, 300, "#F4F8FC", "任务准备"),
        (95, 350, 1605, 525, "#F6FAF4", "报告处理"),
        (95, 575, 1605, 770, "#FBF7F2", "查重复核"),
    ]
    for x1, y1, x2, y2, fill, label in lanes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=fill, outline=(220, 226, 235), width=2)
        d.text((x1 + 24, y1 + 18), label, font=sf, fill=MUTED)
    items = [
        ((250, 205, 500, 270), "教师登录系统"),
        ((730, 205, 980, 270), "创建课程/班级/任务"),
        ((1210, 205, 1460, 270), "选择实验任务"),
        ((250, 410, 500, 485), "上传或批量导入报告"),
        ((730, 410, 980, 485), "解析文本并检测字数/句子数"),
        ((1210, 410, 1460, 485), "保存报告记录"),
        ((250, 645, 500, 720), "发起查重任务"),
        ((730, 645, 980, 720), "生成相似度与风险等级"),
        ((1210, 645, 1460, 720), "查看相似句并导出结果"),
    ]
    for box, txt in items:
        card(d, box, txt, "#FFFFFF", f, outline=(170, 188, 205), radius=14)
    for s, e in [
        ((500, 238), (730, 238)), ((980, 238), (1210, 238)),
        ((1335, 270), (1335, 410)), ((1210, 448), (980, 448)), ((730, 448), (500, 448)),
        ((375, 485), (375, 645)), ((500, 682), (730, 682)), ((980, 682), (1210, 682)),
    ]:
        arrow(d, s, e, width=3)
    text_box(d, "流程终点不是自动判定抄袭，而是为教师提供可追溯的复核依据。", (330, 810, 1370, 875), img_font(26), fill=INK)
    img.save(path)


def make_placeholder(path, cap):
    img = Image.new("RGB", (1350, 540), "#F7F9FB")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((25, 25, 1325, 515), radius=24, fill="#FFFFFF", outline=(170, 182, 196), width=3)
    d.rounded_rectangle((75, 80, 1275, 220), radius=18, fill="#EEF4FA", outline=(190, 204, 220), width=2)
    text_box(d, "系统截图占位", (75, 95, 1275, 145), img_font(34), fill=BLUE)
    text_box(d, "后续将此处替换为实际运行界面截图，保留题注编号即可。", (75, 150, 1275, 205), img_font(24), fill=MUTED)
    text_box(d, cap, (120, 300, 1230, 390), img_font(25), fill=INK)
    img.save(path)


BW_INK = (0, 0, 0)


def bw_title(draw, text, width):
    f = img_font(46)
    b = draw.textbbox((0, 0), text, font=f)
    draw.text(((width - (b[2] - b[0])) / 2, 34), text, font=f, fill=BW_INK)


def bw_text(draw, text, box, font, align="center", spacing=8):
    x1, y1, x2, y2 = box
    max_w = x2 - x1 - 28
    lines = []
    for para in text.split("\n"):
        line = ""
        for ch in para:
            trial = line + ch
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_w:
                line = trial
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
    heights = [draw.textbbox((0, 0), line, font=font)[3] - draw.textbbox((0, 0), line, font=font)[1] for line in lines]
    total_h = sum(heights) + max(0, len(lines) - 1) * spacing
    y = y1 + max(0, (y2 - y1 - total_h) / 2)
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font)[2]
        x = x1 + 16 if align == "left" else x1 + (x2 - x1 - w) / 2
        draw.text((x, y), line, font=font, fill=BW_INK)
        y += h + spacing


def bw_box(draw, box, text, font=None, width=3):
    if font is None:
        font = img_font(28)
    draw.rectangle(box, fill="white", outline=BW_INK, width=width)
    bw_text(draw, text, box, font)


def bw_arrow(draw, start, end, label=None, label_offset=(0, 0), width=4):
    draw.line([start, end], fill=BW_INK, width=width)
    sx, sy = start
    ex, ey = end
    ang = math.atan2(ey - sy, ex - sx)
    size = 20
    pts = [
        (ex, ey),
        (ex - size * math.cos(ang - math.pi / 6), ey - size * math.sin(ang - math.pi / 6)),
        (ex - size * math.cos(ang + math.pi / 6), ey - size * math.sin(ang + math.pi / 6)),
    ]
    draw.polygon(pts, fill=BW_INK)
    if label:
        mx, my = (sx + ex) / 2 + label_offset[0], (sy + ey) / 2 + label_offset[1]
        f = img_font(20)
        b = draw.textbbox((0, 0), label, font=f)
        draw.rectangle((mx - 8, my - 4, mx + b[2] - b[0] + 8, my + b[3] - b[1] + 6), fill="white")
        draw.text((mx, my), label, font=f, fill=BW_INK)


def make_architecture(path):
    img = Image.new("RGB", (2400, 1400), "white")
    d = ImageDraw.Draw(img)
    bw_title(d, "系统总体架构", 2400)
    f, small = img_font(30), img_font(24)
    for x, y, label in [(110, 285, "访问与展示"), (110, 615, "后端业务"), (110, 1010, "数据与算法")]:
        d.text((x, y), label, font=small, fill=BW_INK)
    items = [
        ((300, 245, 680, 345), "浏览器访问\n管理员 / 教师 / 学生"),
        ((840, 245, 1220, 345), "Vue 3 前端\nElement Plus / ECharts"),
        ((1380, 245, 1760, 345), "Axios 请求\nHTTP / JSON"),
        ((300, 575, 680, 695), "Spring Boot API\nController"),
        ((840, 575, 1220, 695), "业务服务层\nService / Mapper"),
        ((1380, 575, 1760, 695), "权限控制\nJWT / 角色校验"),
        ((300, 970, 680, 1090), "MySQL 数据库\n用户 / 任务 / 结果"),
        ((840, 970, 1220, 1090), "文件存储\nuploads 目录"),
        ((1380, 970, 1760, 1090), "查重算法\nTF-IDF / SimHash"),
    ]
    for box, text in items:
        bw_box(d, box, text, f)
    for s, e in [
        ((680, 295), (840, 295)),
        ((1220, 295), (1380, 295)),
        ((1570, 345), (1570, 575)),
        ((1380, 635), (1220, 635)),
        ((840, 635), (680, 635)),
        ((490, 695), (490, 970)),
        ((1030, 695), (1030, 970)),
        ((1570, 695), (1570, 970)),
    ]:
        bw_arrow(d, s, e)
    img.save(path)


def make_function(path):
    img = Image.new("RGB", (2400, 1500), "white")
    d = ImageDraw.Draw(img)
    bw_title(d, "系统功能结构", 2400)
    f, small = img_font(28), img_font(24)
    bw_box(d, (830, 130, 1570, 230), "实验报告智能查重与分析系统", f, width=4)
    columns = [
        (190, "管理员端", ["用户管理", "课程管理", "班级管理", "班级学生分配", "系统看板"]),
        (720, "教师端", ["实验任务管理", "报告上传与批量导入", "发起查重任务", "查看查重结果", "统计分析与导出"]),
        (1250, "学生端", ["查看实验任务", "提交实验报告", "查看提交状态"]),
        (1780, "算法与数据", ["文本解析", "文本预处理", "TF-IDF 余弦相似度", "SimHash 指纹", "风险等级与相似句"]),
    ]
    for x, head, items in columns:
        bw_arrow(d, (1200, 230), (x + 190, 360))
        bw_box(d, (x, 360, x + 380, 450), head, f)
        y = 520
        for item in items:
            bw_box(d, (x, y, x + 380, y + 72), item, small, width=2)
            y += 96
    img.save(path)


def make_algorithm(path):
    img = Image.new("RGB", (2500, 1300), "white")
    d = ImageDraw.Draw(img)
    bw_title(d, "查重算法流程", 2500)
    f, small = img_font(25), img_font(28)
    steps = [
        ("报告集合", "读取同一实验任务下\n全部已解析报告"),
        ("文本预处理", "清洗文本、二元词条切分\n停用词过滤"),
        ("权重建模", "统计文档频率\n计算 TF-IDF 向量"),
        ("相似度计算", "两两计算余弦相似度\n生成 SimHash 指纹"),
        ("融合判定", "0.7×余弦 + 0.3×SimHash\n划分风险等级"),
        ("结果解释", "保存查重结果\n提取相似句详情"),
    ]
    x0, y0, w, h, gap = 110, 300, 330, 160, 70
    for i, (head, body) in enumerate(steps):
        x = x0 + i * (w + gap)
        bw_box(d, (x, y0, x + w, y0 + h), f"{head}\n{body}", f)
        if i < len(steps) - 1:
            bw_arrow(d, (x + w, y0 + h / 2), (x + w + gap, y0 + h / 2))
    d.rectangle((410, 760, 2090, 1000), outline=BW_INK, width=3)
    bw_text(d, "输出结果作为教师复核依据，系统保存综合相似度、风险等级和相似句详情，不直接替代教师判断。", (460, 805, 2040, 955), small)
    img.save(path)


def make_er(path):
    img = Image.new("RGB", (2500, 1320), "white")
    d = ImageDraw.Draw(img)
    bw_title(d, "数据库核心实体关系图", 2500)
    ef, rf, af, small = img_font(28), img_font(22), img_font(18), img_font(20)

    entities = {
        "role": ((240, 260, 460, 340), "角色"),
        "user": ((720, 260, 940, 340), "用户"),
        "class": ((720, 600, 940, 680), "班级"),
        "course": ((240, 900, 460, 980), "课程"),
        "task": ((1060, 900, 1320, 980), "实验任务"),
        "report": ((1640, 620, 1900, 700), "报告文件"),
        "result": ((1640, 1020, 1900, 1100), "查重结果"),
    }

    def rect(box, text):
        d.rectangle(box, fill="white", outline=BW_INK, width=3)
        bw_text(d, text, box, ef)

    def diamond(cx, cy, w, h, text):
        pts = [(cx, cy - h / 2), (cx + w / 2, cy), (cx, cy + h / 2), (cx - w / 2, cy)]
        d.polygon(pts, fill="white", outline=BW_INK)
        d.line([pts[0], pts[1], pts[2], pts[3], pts[0]], fill=BW_INK, width=3)
        bw_text(d, text, (cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2), rf)
        return (cx, cy)

    def ellipse(cx, cy, w, h, text):
        d.ellipse((cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2), fill="white", outline=BW_INK, width=2)
        bw_text(d, text, (cx - w / 2, cy - h / 2, cx + w / 2, cy + h / 2), af)
        return (cx, cy)

    def center(key):
        x1, y1, x2, y2 = entities[key][0]
        return ((x1 + x2) / 2, (y1 + y2) / 2)

    def draw_line(a, b):
        d.line([a, b], fill=BW_INK, width=3)

    def label(text, x, y):
        bw_text(d, text, (x - 40, y - 20, x + 40, y + 20), img_font(19))

    def relation_line(a, b, relation, left_label=None, right_label=None, left_pos=None, right_pos=None):
        draw_line(a, relation)
        draw_line(relation, b)
        if left_label:
            label(left_label, *left_pos)
        if right_label:
            label(right_label, *right_pos)

    relations = {
        "belong": (600, 300, 150, 80, "属于"),
        "join": (830, 465, 150, 80, "加入"),
        "teach": (760, 940, 150, 80, "开设"),
        "arrange": (1015, 740, 150, 80, "安排"),
        "submit": (1335, 500, 150, 80, "提交"),
        "include": (1485, 840, 150, 80, "包含"),
        "compare": (1770, 860, 150, 80, "比较"),
    }

    attrs = [
        ("role", 350, 160, "角色名称"),
        ("user", 830, 150, "用户编号"), ("user", 640, 215, "姓名"), ("user", 1030, 215, "账号"),
        ("class", 615, 570, "班级名称"), ("class", 1030, 645, "专业"),
        ("course", 350, 1085, "课程名称"),
        ("task", 1190, 820, "任务标题"), ("task", 1320, 1085, "截止时间"),
        ("report", 1770, 510, "文件名"), ("report", 2070, 660, "解析文本"),
        ("result", 1770, 1210, "综合相似度"), ("result", 2070, 1060, "风险等级"),
    ]

    relation_line(center("role"), center("user"), (600, 300), "1", "N", (505, 280), (690, 280))
    relation_line(center("user"), center("class"), (830, 465), "M", "N", (855, 380), (855, 555))
    relation_line(center("course"), center("task"), (760, 940), "1", "N", (525, 920), (970, 920))
    relation_line(center("class"), center("task"), (1015, 740), "1", "N", (940, 650), (1070, 850))
    relation_line(center("user"), center("report"), (1335, 500), "1", "N", (1030, 395), (1540, 560))
    relation_line(center("task"), center("report"), (1485, 840), "1", "N", (1380, 880), (1590, 775))
    relation_line(center("report"), center("result"), (1770, 860), "M", "N", (1815, 780), (1815, 970))

    for key, x, y, _ in attrs:
        draw_line(center(key), (x, y))

    for cx, cy, w, h, text in relations.values():
        diamond(cx, cy, w, h, text)

    for key, x, y, text in attrs:
        ellipse(x, y, 135, 50, text)

    for box, text in entities.values():
        rect(box, text)

    img.save(path)


def make_tech_flow(path):
    img = Image.new("RGB", (2500, 1400), "white")
    d = ImageDraw.Draw(img)
    bw_title(d, "系统技术流程", 2500)
    f, label_f = img_font(26), img_font(24)
    for x, y, label in [(150, 300, "提交阶段"), (150, 650, "处理阶段"), (150, 1000, "查重阶段")]:
        d.text((x, y), label, font=label_f, fill=BW_INK)
    items = [
        ((330, 285, 690, 365), "选择实验任务"),
        ((860, 285, 1220, 365), "上传报告文件"),
        ((1390, 285, 1750, 365), "前端提交请求"),
        ((330, 620, 690, 710), "接口鉴权与参数校验"),
        ((860, 620, 1220, 710), "保存文件并解析文本"),
        ((1390, 620, 1750, 710), "保存报告记录"),
        ((330, 970, 690, 1060), "创建查重任务"),
        ((860, 970, 1220, 1060), "计算相似度"),
        ((1390, 970, 1750, 1060), "写入结果与相似句"),
        ((1920, 970, 2280, 1060), "查看结果与图表"),
    ]
    for box, text in items:
        bw_box(d, box, text, f, width=2)
    for s, e in [
        ((690, 325), (860, 325)),
        ((1220, 325), (1390, 325)),
        ((1570, 365), (1570, 620)),
        ((1390, 665), (1220, 665)),
        ((860, 665), (690, 665)),
        ((510, 710), (510, 970)),
        ((690, 1015), (860, 1015)),
        ((1220, 1015), (1390, 1015)),
        ((1750, 1015), (1920, 1015)),
    ]:
        bw_arrow(d, s, e, width=3)
    img.save(path)


def make_progress(path):
    img = Image.new("RGB", (2400, 1300), "white")
    d = ImageDraw.Draw(img)
    bw_title(d, "实验进度与成果", 2400)
    f = img_font(25)
    stages = [
        ("阶段一\n需求与开题", "完成开题报告\n确定技术路线"),
        ("阶段二\n基础框架", "后端 / 前端 / 数据库\n完成基础搭建"),
        ("阶段三\n核心功能", "报告上传解析\n查重算法实现"),
        ("阶段四\n演示数据", "多班级多任务样例\n批量上传验证"),
        ("阶段五\n测试与论文", "单元测试 / 构建验证\n整理报告终稿"),
    ]
    x, y, w, h, gap = 170, 420, 330, 170, 105
    d.line((260, y + h + 95, 2140, y + h + 95), fill=BW_INK, width=5)
    for i, (head, body) in enumerate(stages):
        bx = x + i * (w + gap)
        bw_box(d, (bx, y, bx + w, y + h), f"{head}\n{body}", f)
        cx = bx + w / 2
        d.ellipse((cx - 24, y + h + 71, cx + 24, y + h + 119), fill="white", outline=BW_INK, width=4)
        d.text((cx - 8, y + h + 74), str(i + 1), font=img_font(25), fill=BW_INK)
    d.rectangle((390, 910, 2010, 1060), outline=BW_INK, width=3)
    bw_text(d, "当前系统已完成可运行闭环：登录 → 任务 → 上传 → 解析 → 查重 → 结果 → 图表 → 导出。", (430, 940, 1970, 1030), img_font(30))
    img.save(path)


def make_business(path):
    img = Image.new("RGB", (2400, 1350), "white")
    d = ImageDraw.Draw(img)
    bw_title(d, "教师端查重业务流程", 2400)
    f, label_f = img_font(26), img_font(24)
    for x, y, label in [(130, 275, "任务准备"), (130, 625, "报告处理"), (130, 1000, "查重复核")]:
        d.text((x, y), label, font=label_f, fill=BW_INK)
    items = [
        ((300, 255, 650, 330), "教师登录系统"),
        ((1025, 255, 1375, 330), "创建课程 / 班级 / 任务"),
        ((1750, 255, 2100, 330), "选择实验任务"),
        ((300, 595, 650, 690), "上传或批量导入报告"),
        ((1025, 595, 1375, 690), "解析文本并检测字数 / 句子数"),
        ((1750, 595, 2100, 690), "保存报告记录"),
        ((300, 970, 650, 1065), "发起查重任务"),
        ((1025, 970, 1375, 1065), "生成相似度与风险等级"),
        ((1750, 970, 2100, 1065), "查看相似句并导出结果"),
    ]
    for box, text in items:
        bw_box(d, box, text, f, width=2)
    for s, e in [((650, 292), (1025, 292)), ((1375, 292), (1750, 292)), ((1925, 330), (1925, 595)), ((1750, 642), (1375, 642)), ((1025, 642), (650, 642)), ((475, 690), (475, 970)), ((650, 1017), (1025, 1017)), ((1375, 1017), (1750, 1017))]:
        bw_arrow(d, s, e, width=3)
    d.rectangle((430, 1180, 1970, 1260), outline=BW_INK, width=3)
    bw_text(d, "流程终点是为教师提供可追溯的复核依据，而不是自动判定学生抄袭。", (470, 1195, 1930, 1248), img_font(28))
    img.save(path)


def make_placeholder(path, cap):
    img = Image.new("RGB", (1800, 720), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((35, 35, 1765, 685), outline=BW_INK, width=4)
    d.rectangle((110, 105, 1690, 260), outline=BW_INK, width=3)
    bw_text(d, "系统截图占位", (110, 122, 1690, 178), img_font(38))
    bw_text(d, "后续将此处替换为实际运行界面截图，保留题注编号即可。", (110, 180, 1690, 238), img_font(26))
    bw_text(d, cap, (170, 390, 1630, 500), img_font(28))
    img.save(path)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(*(color or (0, 0, 0)))


def set_para(p, line=1.5, before=0, after=0, first=True, align=None):
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Cm(0.74) if first else Cm(0)
    if align is not None:
        p.alignment = align


def add_p(doc, text="", first=True):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        set_run_font(r, 12)
    set_para(p, first=first)
    return p


def title_p(doc, text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size, bold=True)
    set_para(p, after=10, first=False)
    return p


def heading(doc, text, level=1, center=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 16 if level == 1 else 14 if level == 2 else 12, bold=True)
    set_para(p, before=8 if level == 1 else 4, after=4, first=False)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 10.5)
    set_para(p, line=1.2, after=6, first=False)


def add_pic(doc, path, cap):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(14.5))
    set_para(p, line=1.0, after=2, first=False)
    caption(doc, cap)


def set_cell(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(text)) > 12 else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_run_font(r, 10.5, bold=bold)
    set_para(p, line=1.15, first=False)


def shade(cell, fill):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc.append(shd)


def table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "888888")


def add_table(doc, headers, rows, cap):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    table_borders(table)
    caption(doc, cap)
    return table


def setup_section(sec):
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.75)
    sec._sectPr.pgMar.set(qn("w:gutter"), str(int(Cm(1.0).twips)))


def add_header_footer(sec, roman=False):
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    for p in sec.header.paragraphs:
        p.clear()
    for p in sec.footer.paragraphs:
        p.clear()
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(TITLE)
    set_run_font(r, 10.5)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    rr = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "I" if roman else "1"
    rr.append(t)
    fld.append(rr)
    fp._p.append(fld)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), "upperRoman" if roman else "decimal")
    pg.set(qn("w:start"), "1")
    sec._sectPr.append(pg)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for tag, attrs, text in [
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, 'TOC \\o "1-3" \\h \\z \\u'),
        ("w:fldChar", {"w:fldCharType": "separate"}, None),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ]:
        node = OxmlElement(tag)
        for k, v in attrs.items():
            node.set(qn(k), v)
        if text:
            node.text = text
        run._r.append(node)
    set_para(p, first=False)


def build_base_doc(out, abstract_cn, abstract_en):
    doc = Document()
    setup_section(doc.sections[0])
    doc.sections[0].header.is_linked_to_previous = False
    doc.sections[0].footer.is_linked_to_previous = False
    for s in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[s]
        st.font.name = "Times New Roman"
        st.font.color.rgb = RGBColor(0, 0, 0)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    doc.styles["Normal"].font.size = Pt(12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("评阅教师：____________")
    set_run_font(r, 12)
    set_para(p, first=False)
    doc.add_paragraph()
    title_p(doc, "海南大学计算机科学与技术学院", 16)
    title_p(doc, "计算机综合课程设计报告", 22)
    for _ in range(3):
        doc.add_paragraph()
    title_p(doc, TITLE, 18)
    for label, value in [("班    级", "（请填写）"), ("姓    名", "（请填写）"), ("学    号", "（请填写）"), ("组    员", "无"), ("指导老师", "（请填写）"), ("完成日期", "2026年6月")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}：        {value}        ")
        set_run_font(r, 14)
        set_para(p, line=1.8, first=False)

    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_section(sec)
    add_header_footer(sec, roman=True)
    title_p(doc, "摘    要", 16)
    add_p(doc, abstract_cn)
    p = doc.add_paragraph()
    r = p.add_run("关键词：")
    set_run_font(r, 12, bold=True)
    r = p.add_run("实验报告查重；文本相似度；TF-IDF；余弦相似度；SimHash")
    set_run_font(r, 12)
    set_para(p, first=False)
    doc.add_page_break()
    title_p(doc, "Abstract", 16)
    add_p(doc, abstract_en)
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    set_run_font(r, 12, bold=True)
    r = p.add_run("Experimental report checking; Text similarity; TF-IDF; Cosine similarity; SimHash")
    set_run_font(r, 12)
    set_para(p, first=False)
    doc.add_page_break()
    title_p(doc, "目    录", 16)
    add_toc(doc)
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_section(sec)
    add_header_footer(sec)
    return doc


def add_common_references(doc):
    doc.add_page_break()
    heading(doc, "致    谢", 1, center=True)
    add_p(doc, "在本次计算机综合课程设计过程中，我围绕实验报告查重这一教学场景完成了需求分析、系统设计、数据库建模、前后端开发、查重算法实现、测试验证和文档整理等工作。通过该项目，我进一步理解了 Web 系统开发流程，也加深了对文本相似度算法和工程实践的认识。")
    add_p(doc, "感谢指导教师在课程设计选题、技术路线和报告撰写方面给予的指导，也感谢课程学习过程中提供帮助的同学和相关开源技术社区。由于个人能力和时间有限，系统仍存在不足，后续将继续从算法效果、系统安全性和用户体验等方面进行完善。")
    doc.add_page_break()
    heading(doc, "参考文献", 1, center=True)
    refs = [
        "[1] Salton G, Buckley C. Term-weighting approaches in automatic text retrieval[J]. Information Processing & Management, 1988, 24(5): 513-523.",
        "[2] Manku G S, Jain A, Das Sarma A. Detecting near-duplicates for web crawling[C]. Proceedings of the 16th International Conference on World Wide Web, 2007: 141-150.",
        "[3] 李智慧. Spring Boot 企业级开发教程[M]. 北京: 人民邮电出版社, 2021.",
        "[4] 尤雨溪. Vue.js 设计与实现[M]. 北京: 人民邮电出版社, 2022.",
        "[5] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "[6] 宗成庆. 统计自然语言处理[M]. 北京: 清华大学出版社, 2013.",
        "[7] Halim J, Lasut D. Document Plagiarism Detection Application Using Web-Based TF-IDF and Cosine Similarity Methods[J]. bit-Tech, 2024, 7(2): 202-213.",
        "[8] Li Z, Zhu J, Cheng X, Lu Q. Optimized Algorithm Design for Text Similarity Detection Based on Artificial Intelligence and Natural Language Processing[J]. Procedia Computer Science, 2023, 228: 195-202.",
        "[9] Yuan L, Gao S, Pan P. CTSARF: A Chinese Text Similarity Analysis Model based on Residual Fusion[J]. Neurocomputing, 2023, 559: 126801.",
        "[10] Saeed S, Rajput Q, Haider S. SUMEX: A hybrid framework for Semantic textUal siMilarity and EXplanation generation[J]. Information Processing & Management, 2024, 61(5): 103771.",
        "[11] Apache Software Foundation. Apache POI Documentation[EB/OL]. https://poi.apache.org/.",
        "[12] MyBatis-Plus Team. MyBatis-Plus Documentation[EB/OL]. https://baomidou.com/.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        set_run_font(r, 10.5)
        set_para(p, line=1.2, first=False)


def build_initial():
    assets = {
        "algorithm": ASSET_DIR / "initial_algorithm.png",
        "tech": ASSET_DIR / "initial_tech_flow.png",
        "progress": ASSET_DIR / "initial_progress.png",
        "architecture": ASSET_DIR / "initial_architecture.png",
    }
    make_algorithm(assets["algorithm"])
    make_tech_flow(assets["tech"])
    make_progress(assets["progress"])
    make_architecture(assets["architecture"])
    doc = build_base_doc(
        INITIAL_OUT,
        "本文围绕高校实验报告批量复核场景，设计并实现基于 Spring Boot 与文本相似度算法的实验报告查重系统初稿。系统采用前后端分离架构，后端基于 Spring Boot、MyBatis-Plus 和 MySQL 完成登录认证、课程班级管理、实验任务管理、报告上传、文本解析和查重计算，前端基于 Vue 3、Element Plus 和 ECharts 完成后台管理界面与统计展示。本文初稿重点讨论课题实验设置、实验数据、实验进度和阶段性成果，并给出后续实验问题的解决方案。算法方面，系统以同一实验任务下全部已解析报告作为语料集合，采用轻量级中文二元词条切分、TF-IDF 权重、余弦相似度和 SimHash 指纹相似度计算综合相似度，并按照阈值划分风险等级。当前系统已经完成可运行闭环，具备继续完善和撰写终稿的基础。",
        "This draft focuses on the design and implementation progress of an experimental report plagiarism checking system based on Spring Boot and text similarity algorithms. The system adopts a front-end and back-end separated architecture. The back end uses Spring Boot, MyBatis-Plus and MySQL to implement authentication, course and class management, experiment task management, report upload, text extraction and similarity calculation. The front end uses Vue 3, Element Plus and ECharts to implement management pages and data visualization. This draft mainly discusses experimental settings, experimental data, development progress, preliminary results and solutions to specific problems. In terms of algorithms, all parsed reports under the same task are used as the corpus. Lightweight Chinese bigram tokenization, TF-IDF weighting, cosine similarity and SimHash fingerprint similarity are combined to calculate final similarity and risk levels. The current system has completed a runnable workflow and provides a feasible basis for the final report.",
    )
    heading(doc, "1 绪论", 1)
    add_p(doc, "实验报告查重系统的研究对象不是普通长篇论文，而是高校实验教学中按课程、班级和实验任务集中提交的短文本报告。该类报告通常具有固定模板、相似实验步骤和较多专业术语，人工复核时容易受到工作量和主观经验影响。因此，本课题将系统目标定位为构建一个轻量级、可解释、可演示的辅助查重系统，为教师提供相似度、风险等级和相似句等复核依据。")
    add_p(doc, "初稿阶段的主要任务不是追求系统的产品化完备性，而是论证实验方法是否可行、数据与流程是否完整、核心算法能否产生可解释结果。围绕这一目标，本文先说明课题实验设置和数据来源，再给出技术流程与算法流程，最后总结当前进度、阶段成果和待解决问题。")
    heading(doc, "2 课题实验设置", 1)
    add_p(doc, "本章围绕课题实验如何开展进行说明。实验设置需要回答三个问题：系统要验证什么、使用哪些数据验证、业务流程如何串联到算法流程。只有先明确实验目标、实验数据和技术流程，后续的算法设计与系统测试才具有可复现依据。")
    heading(doc, "2.1 实验目标", 2)
    add_p(doc, "本课题实验目标包括三个层面。第一，验证系统业务流程可行，即教师能够创建实验任务、上传多份报告、发起查重并查看结果；第二，验证文本处理和相似度算法可行，即系统能够从 txt、docx、pdf 文件中提取文本，并计算不同报告之间的相似度；第三，验证结果表达可用，即系统输出不仅包含数值，还包含风险等级、相似句和统计图表，便于教师进行人工复核。")
    heading(doc, "2.2 实验数据", 2)
    add_p(doc, "为避免系统只停留在单个样例演示，初稿阶段构造了多班级、多学生和多任务的数据环境。初始化脚本中包含管理员、教师、学生、课程、班级和实验任务等基础数据，样例报告则覆盖完全重复、局部改写、主题不同和接口开发报告相似等场景。为了让这些数据与实验目标之间的对应关系更加清楚，表2-1 将数据类型、规模和验证用途进行归纳。")
    add_table(doc, ["数据类型", "规模", "用途"], [["用户数据", "1 个管理员、2 个教师、8 个学生", "验证角色权限和多学生报告提交"], ["班级数据", "3 个班级", "验证多班级管理"], ["课程数据", "3 门课程", "验证课程与教师关系"], ["实验任务", "4 个任务", "验证不同任务下独立查重"], ["样例报告", "多组 txt 报告", "验证高风险、中风险、低风险和正常结果"]], "表2-1 初稿阶段实验数据设置")
    add_p(doc, "从表2-1 可以看出，实验数据并不是简单凑数量，而是围绕系统主流程逐项设置。用户、班级和课程数据用于验证权限与组织关系，实验任务用于限定查重范围，样例报告则用于观察不同相似程度下的风险等级输出。这样的数据设置能够支撑登录、任务、上传、解析、查重和结果展示的连续演示。")
    heading(doc, "2.3 技术流程", 2)
    add_p(doc, "为了保证实验流程能够从用户操作延伸到后端算法处理，系统将前端交互、后端业务、数据库存储和算法计算串联起来。图2-1 按提交阶段、处理阶段和查重阶段展示系统从上传报告到查看结果的技术流程，用于说明实验并不是孤立运行算法，而是在 Web 系统中完成完整业务闭环。")
    add_pic(doc, assets["tech"], "图2-1 系统技术流程图")
    add_p(doc, "图2-1 中，上传报告首先经过前端提交和接口鉴权，然后进入文件保存、文本解析和报告记录保存，最后由查重任务触发相似度计算并写入结果。该流程说明系统不是单独调用算法函数，而是在真实 Web 业务中完成文件上传、权限校验、文本解析、数据保存和结果展示，这使实验结果更接近课程设计实际要求。")
    heading(doc, "3 初步实验方法", 1)
    add_p(doc, "本章说明初稿阶段已经采用的实验方法。系统先完成报告文件的格式检查和文本解析，再对同一实验任务下的报告进行预处理、向量化和相似度计算，最终输出综合相似度和风险等级。该方法强调可实现性和可解释性，适合课程设计阶段验证系统主流程。")
    heading(doc, "3.1 文本解析方法", 2)
    add_p(doc, "报告上传后，系统首先判断文件是否为空、大小是否超过 20MB、文件类型是否为 txt、docx 或 pdf。对于 txt 文件，系统优先使用 UTF-8 编码读取，失败后兼容 GBK；对于 docx 文件，系统使用 Apache POI 提取段落文本；对于 pdf 文件，系统使用 PDFBox 提取普通文本内容。解析成功后，系统记录文本字数、句子数和解析状态。")
    heading(doc, "3.2 相似度计算方法", 2)
    add_p(doc, "查重实验以同一实验任务作为比较范围。系统读取该任务下所有解析成功的报告，完成文本清洗、中文二元词条切分和停用词过滤，然后基于报告集合统计文档频率并计算 TF-IDF 权重。每两份报告之间先计算余弦相似度，再计算 SimHash 指纹相似度，最终按 0.7 和 0.3 的权重融合。")
    add_p(doc, "为了把上述文字处理过程与代码实现对应起来，图3-1 将报告集合读取、文本预处理、权重建模、相似度计算、风险判定和结果解释串联为一条算法流程。")
    add_pic(doc, assets["algorithm"], "图3-1 查重算法流程图")
    add_p(doc, "图3-1 中最重要的是将报告集合限定在同一实验任务内，因为不同任务的报告主题差异较大，跨任务比较会降低结果解释性。当前算法更适合识别完全复制、大段复制和轻度改写，对于深度语义改写仍需后续算法增强，因此初稿中需要把算法定位为可解释的辅助复核方法，而不是深度语义查重模型。")
    heading(doc, "4 实验进度与阶段成果", 1)
    add_p(doc, "目前课题已经完成开题报告、数据库建模、后端基础接口、前端后台页面、报告上传解析、查重算法、查重结果展示、统计图表和演示数据构造等工作。为了让阶段成果与后续任务之间的关系更加直观，图4-1 将当前实验进度按五个阶段归纳。")
    add_pic(doc, assets["progress"], "图4-1 实验进度与阶段成果图")
    add_p(doc, "从图4-1 可以看出，课题已经完成需求与开题、基础框架搭建、核心功能开发、演示数据准备和论文材料整理等阶段。当前阶段成果不是单独的算法片段，而是已经形成登录、任务、上传、解析、查重、结果、图表和导出的完整流程。后续工作主要集中在论文表达完善、前端页面结构优化和更多测试用例扩展。")
    heading(doc, "5 实验问题与解决方法", 1)
    add_p(doc, "本章针对初稿阶段已经暴露出的具体问题进行说明。这里的问题并不是系统无法运行的错误，而是课程设计论文需要如实交代的边界条件，包括中文分词精度、实验报告模板相似、上传格式差异、演示数据规模和系统安全性定位等。")
    add_p(doc, "初稿阶段需要重点说明几个实验问题及其处理方式。第一，当前中文文本处理采用轻量级二元词条切分，无法完全替代 Jieba 或 HanLP 等成熟分词工具，因此论文中不应夸大为深度语义理解，而应表述为适合课程设计规模的轻量级预处理方法。第二，实验报告往往存在固定模板，可能导致正常内容相似度偏高，所以系统提供风险等级和相似句详情，由教师结合任务模板进行复核。第三，上传文件格式不统一会影响解析结果，系统通过 txt、docx、pdf 三类格式支持、解析状态记录和失败原因提示来降低演示风险。第四，为避免系统显得只适合单个样例，演示数据增加了多班级、多学生、多任务和多组样例报告。第五，系统安全性定位为基础认证和角色权限控制，不把当前明文密码实现表述为强安全方案，后续可继续引入 BCrypt。")
    add_p(doc, "这些问题并不影响系统完成课程设计核心目标，但需要在论文中如实说明。尤其是算法部分应强调系统定位为辅助复核工具，而不是自动判定抄袭的平台。")
    heading(doc, "6 小结", 1)
    add_p(doc, "初稿阶段已经验证了课题实验方法的可行性。系统具备真实数据、可运行流程和可解释算法结果，能够支撑终稿撰写和答辩演示。下一阶段将重点完善论文表达、优化系统结构说明，并继续补充测试材料。")
    add_common_references(doc)
    doc.add_page_break()
    heading(doc, "附    件", 1, center=True)
    heading(doc, "附件 A 项目材料说明", 2)
    add_p(doc, "项目材料由后端工程、前端工程、数据库脚本、样例报告和说明文档组成。backend 目录保存 Spring Boot 后端代码，frontend 目录保存 Vue 前端代码，database/init.sql 保存建表语句和初始化演示数据，samples 目录保存用于批量上传和查重测试的样例报告。")
    add_p(doc, "这些材料共同支撑本文所述的实验流程。后端工程负责登录认证、基础数据管理、报告上传解析和查重计算，前端工程负责页面交互和结果展示，数据库脚本负责构造用户、课程、班级、实验任务和报告相关数据，样例报告用于验证不同相似程度下的风险等级输出。")
    doc.save(INITIAL_OUT)
    return INITIAL_OUT


def make_er_complete(path):
    img = Image.new("RGB", (1750, 1080), "white")
    d = ImageDraw.Draw(img)
    title(d, "数据库完整 ER 关系图", 1750)
    f, sf = img_font(22), img_font(18)
    nodes = {
        "sys_role\n角色": (120, 170, 350, 270),
        "sys_user\n用户": (500, 170, 730, 270),
        "class_info\n班级": (880, 170, 1110, 270),
        "student_class\n学生班级关联": (1260, 170, 1530, 270),
        "course\n课程": (500, 420, 730, 520),
        "experiment_task\n实验任务": (880, 420, 1150, 520),
        "report_file\n报告文件": (1260, 420, 1530, 520),
        "check_task\n查重任务": (880, 690, 1150, 790),
        "check_result\n查重结果": (1260, 690, 1530, 790),
        "similar_sentence\n相似句子": (1260, 900, 1530, 1000),
    }
    for text, box in nodes.items():
        card(d, box, text, "#F7FAFD", f, outline=(130, 155, 185), radius=14)

    def c(box):
        return ((box[0] + box[2]) / 2, (box[1] + box[3]) / 2)

    relations = [
        ("sys_role\n角色", "sys_user\n用户", "1:N"),
        ("sys_user\n用户", "course\n课程", "教师 1:N"),
        ("sys_user\n用户", "student_class\n学生班级关联", "学生 1:N"),
        ("class_info\n班级", "student_class\n学生班级关联", "1:N"),
        ("course\n课程", "experiment_task\n实验任务", "1:N"),
        ("class_info\n班级", "experiment_task\n实验任务", "1:N"),
        ("experiment_task\n实验任务", "report_file\n报告文件", "1:N"),
        ("sys_user\n用户", "report_file\n报告文件", "学生 1:N"),
        ("experiment_task\n实验任务", "check_task\n查重任务", "1:N"),
        ("check_task\n查重任务", "check_result\n查重结果", "1:N"),
        ("report_file\n报告文件", "check_result\n查重结果", "源/目标报告"),
        ("check_result\n查重结果", "similar_sentence\n相似句子", "1:N"),
    ]
    for a, b, rel in relations:
        start, end = c(nodes[a]), c(nodes[b])
        arrow(d, start, end, width=3)
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        d.text((mx - 42, my - 24), rel, font=sf, fill=MUTED)
    d.text((110, 1015), "说明：report_file 通过 task_id/student_id 限定一名学生在一个实验任务下保留一份正式报告；check_result 同时引用源报告与目标报告。", font=sf, fill=MUTED)
    img.save(path)


def build_final():
    assets = {
        "architecture": ASSET_DIR / "final_architecture.png",
        "function": ASSET_DIR / "final_function.png",
        "algorithm": ASSET_DIR / "final_algorithm.png",
        "er": ASSET_DIR / "final_er.png",
        "tech": ASSET_DIR / "final_tech_flow.png",
    }
    make_architecture(assets["architecture"])
    make_function(assets["function"])
    make_algorithm(assets["algorithm"])
    make_er_complete(assets["er"])
    make_tech_flow(assets["tech"])
    doc = build_base_doc(
        FINAL_OUT,
        "随着高校实验教学和课程实践规模不断扩大，实验报告成为评价学生实践能力和课程学习效果的重要依据。传统人工检查方式在面对同一实验任务下的大批量报告时效率较低，且难以及时发现大段复制、模板套用和局部改写等问题。针对这一场景，本文设计并实现了一套基于 Spring Boot 与文本相似度算法的实验报告查重系统。系统采用前后端分离架构，后端基于 Spring Boot、MyBatis-Plus 和 MySQL 实现用户认证、课程管理、班级管理、实验任务管理、报告上传、文本解析、查重计算和结果导出等功能，前端基于 Vue 3、Element Plus 和 ECharts 实现后台管理界面、统计看板和查重结果展示。在算法方面，系统支持 txt、docx 和 pdf 报告文本提取，对文本进行清洗、轻量级中文二元词条切分和停用词过滤，并以同一实验任务下的全部已解析报告作为语料集合计算 TF-IDF 权重。系统通过余弦相似度衡量报告向量之间的相似程度，同时引入 SimHash 文本指纹方法计算近重复文本相似度，最终按照加权融合方式得到综合相似度，并划分正常、低风险、中风险和高风险等级。系统还提供相似句子展示、统计图表、CSV 结果导出和多班级多任务演示数据，能够为教师批量复核实验报告提供可量化、可解释的辅助依据。测试结果表明，系统主流程完整，能够完成从任务创建、报告上传、文本解析、发起查重到结果展示的闭环操作，满足课程设计预期目标。",
        "With the expansion of experimental teaching and practical courses in universities, experimental reports have become important materials for evaluating students' practical ability and learning outcomes. Manual review is inefficient when teachers need to inspect a large number of reports under the same experimental task, and it is difficult to identify large-scale copying, template reuse and partial rewriting in time. To address this problem, this paper designs and implements an experimental report plagiarism checking system based on Spring Boot and text similarity algorithms. The system adopts a front-end and back-end separated architecture. The back end is built with Spring Boot, MyBatis-Plus and MySQL, implementing user authentication, course management, class management, experiment task management, report upload, text extraction, similarity calculation and result export. The front end is built with Vue 3, Element Plus and ECharts, providing management pages, dashboards and result visualization. In terms of algorithms, the system extracts text from txt, docx and pdf files, performs text cleaning, lightweight Chinese bigram tokenization and stop-word filtering, and calculates TF-IDF weights based on all parsed reports under the same task. Cosine similarity is used to measure vector similarity, while SimHash is introduced to evaluate near-duplicate text fingerprints. The final similarity score is obtained through weighted fusion and mapped to normal, low-risk, medium-risk and high-risk levels. The system also provides similar sentence display, statistical charts, CSV export and multi-class demonstration data, offering quantitative and interpretable support for teachers. Test results show that the system can complete the full workflow from task creation, report upload and text parsing to plagiarism checking and result display, meeting the expected objectives of the course project.",
    )
    # concise but complete final body; keeps figures tightly integrated
    for chap, paras in [
        ("1 绪论", ["随着高校实验教学、课程设计和综合实践教学规模不断扩大，实验报告已经成为记录学生实践过程、反映课程理解程度和评价动手能力的重要材料。传统人工检查方式依赖教师经验，面对批量报告时效率较低，难以及时定位疑似重复内容。通用论文查重系统通常面向毕业论文和期刊论文，使用成本较高，也不一定适合实验报告这种篇幅较短、结构固定、按课程任务提交的教学场景。因此，本课题设计面向高校实验教学的轻量级实验报告查重系统，为教师提供可量化、可复核的辅助依据。", "本文主要完成系统需求分析、总体架构设计、数据库设计、查重算法设计、核心功能实现和系统测试。系统不是单纯的文本比较工具，而是围绕课程、班级、实验任务、报告文件和查重结果形成完整业务链路。"]),
        ("2 理论基础与关键技术", ["系统采用 B/S 架构和前后端分离模式，前端使用 Vue 3、Element Plus 和 ECharts，后端使用 Spring Boot、MyBatis-Plus 和 MySQL。文本相似度算法方面，系统采用 TF-IDF 权重、余弦相似度和 SimHash 文本指纹。TF-IDF 用于衡量词项在报告集合中的重要程度，余弦相似度用于衡量文本向量夹角，SimHash 则通过文本指纹辅助判断近重复文本。", "考虑到课程设计周期和部署成本，系统没有采用复杂深度语义模型，而是采用轻量级中文二元词条切分方法。这种方法无法完全解决深度语义改写问题，但具有实现简单、计算成本低、结果可解释的特点，适合实验报告近重复检测场景。"]),
        ("3 需求分析", ["系统用户包括管理员、教师和学生。管理员维护用户、课程、班级和班级学生关联；教师创建实验任务、上传或批量导入报告、发起查重并查看结果；学生查看实验任务并提交本人报告。为了让需求和后续测试具有对应关系，系统将功能拆分为登录认证、基础数据管理、实验任务管理、报告上传解析、查重计算、结果展示、统计分析和结果导出等模块。"]),
    ]:
        heading(doc, chap, 1)
        for p in paras:
            add_p(doc, p)
        if chap.startswith("3 "):
            add_p(doc, "开题报告中的功能承诺已经落实到登录认证、基础数据管理、报告上传解析、查重计算、结果展示、统计分析和结果导出等模块中。后续的总体设计、数据库设计和功能测试均围绕这些模块展开。")

    heading(doc, "4 系统总体设计", 1)
    add_p(doc, "系统采用前后端分离架构。为了说明浏览器、前端工程、后端服务、数据库、文件存储和算法模块之间的关系，本文将系统整体结构整理为图4-1，使各层之间的数据流和调用边界更加直观。")
    add_pic(doc, assets["architecture"], "图4-1 系统总体架构图")
    add_p(doc, "图4-1 表明前端并不直接访问数据库和文件系统，而是通过后端 API 完成业务操作；算法模块也位于后端服务内部，便于统一读取报告文本并保存查重结果。这一结构既符合前后端分离开发方式，也便于后续扩展新的文本解析器或相似度算法。")
    add_p(doc, "在功能层面，系统围绕管理员、教师、学生和算法数据支撑四个部分展开。图4-2 将功能边界与角色职责联系起来，用于说明系统不是孤立的算法程序，而是面向实验教学场景的后台管理系统。")
    add_pic(doc, assets["function"], "图4-2 系统功能结构图")
    add_p(doc, "教师端是系统演示主线，贯穿任务创建、报告上传、发起查重、查看结果和导出文件等流程。为了说明这些操作如何在技术层面串联，图4-3 给出了系统技术流程，展示从前端提交到后端保存、算法计算和结果展示的衔接关系。")
    add_pic(doc, assets["tech"], "图4-3 系统技术流程图")

    heading(doc, "5 数据库设计", 1)
    add_p(doc, "数据库设计围绕实验报告查重业务展开，主要保存用户、角色、课程、班级、实验任务、报告文件、查重任务、查重结果和相似句子等信息。由于查重结果需要回溯到具体任务、学生和报告文件，仅用文字描述表关系不够直观，因此本文将核心关系整理为图5-1，并在图中保留一对多、多对多等主要基数关系。")
    add_pic(doc, assets["er"], "图5-1 数据库核心关系图")
    add_p(doc, "图5-1 主要回答实体之间如何关联的问题，而表5-1 进一步说明这些实体在代码落库时分别承担什么职责。")
    add_table(doc, ["数据表", "作用"], [["sys_user / sys_role", "保存用户与角色信息"], ["course / class_info / student_class", "保存课程、班级和学生归属"], ["experiment_task", "保存实验任务"], ["report_file", "保存报告文件与解析文本"], ["check_task / check_result", "保存查重任务与两两比较结果"], ["similar_sentence", "保存相似句详情"]], "表5-1 核心数据表作用")
    add_p(doc, "结合图5-1 和表5-1 可以看出，系统能够从查重结果继续追溯到任务、报告、学生和相似句，为教师复核提供数据依据，而不是只保存一个无法解释的相似度数值。")

    heading(doc, "6 查重算法设计与实现", 1)
    add_p(doc, "系统首先对 txt、docx 和 pdf 报告进行文本提取，然后进行文本清洗、中文二元词条切分和停用词过滤。查重时，系统读取同一实验任务下所有解析成功报告，统一统计文档频率并计算 TF-IDF 权重，再对报告两两计算余弦相似度和 SimHash 相似度。")
    add_p(doc, "为了让算法流程与代码实现形成对应，图6-1 将报告读取、预处理、权重建模、相似度计算、风险判定和相似句保存串联起来，展示系统如何从报告集合逐步生成可解释的查重结果。")
    add_pic(doc, assets["algorithm"], "图6-1 查重算法流程图")
    add_p(doc, "图6-1 中的综合相似度需要进一步转化为教师容易理解的结果，因此表6-1 将分值区间映射为四类风险等级。")
    add_table(doc, ["综合相似度", "风险等级", "处理建议"], [[">=0.80", "高风险", "重点复核"], ["0.60-0.80", "中风险", "查看相似句"], ["0.40-0.60", "低风险", "结合模板判断"], ["<0.40", "正常", "一般检查"]], "表6-1 风险等级划分")
    add_p(doc, "风险等级并不是为了替代教师判断，而是降低教师筛选结果的成本。教师可以优先查看高风险和中风险结果，再结合相似句内容判断相似来源是否属于异常重复。")

    heading(doc, "7 系统功能实现", 1)
    impl = [
        ("7.1 登录认证与权限控制", "登录模块校验用户名、密码和用户状态，成功后生成 JWT。系统同时提供教师和学生注册、绑定邮箱找回密码等入口，管理员账号不允许前台注册。后端拦截器解析 Token 并识别角色，学生只能查看和提交自己的报告，教师和管理员则根据角色访问对应功能。"),
        ("7.2 报告上传与内容检测", "报告上传模块支持单文件上传和批量上传。系统在保存文件前检查文件为空、大小、类型、任务状态、学生身份以及学生是否属于当前实验任务班级；解析成功后记录字数和句子数量，文本偏短时提示教师复核。"),
        ("7.3 查重结果与相似句展示", "查重结果列表按照综合相似度排序，展示报告组合、文件名、余弦相似度、SimHash 相似度、综合相似度和风险等级。系统同步写入 similar_sentence 表，详情弹窗展示相似自然句；若没有相似句，也会显示中文原因说明。"),
        ("7.4 统计分析与结果导出", "统计分析页面通过图表展示风险等级占比、核心数据统计和相似度区间分布。系统支持 CSV 导出，便于教师留存查重结果。"),
    ]
    for sub, para in impl:
        heading(doc, sub, 2)
        add_p(doc, para)

    heading(doc, "8 系统测试与结果分析", 1)
    add_p(doc, "测试工作围绕系统主流程、算法工具类和构建结果展开。功能测试验证登录、注册、密码找回、角色菜单、课程班级管理、实验任务、txt/docx/pdf 报告上传、批量上传、查重任务、相似句详情、删除失败中文提示和结果导出；自动化测试验证相似度计算、空文本处理、风险等级阈值、相似句匹配以及 UTF-8/GBK/docx/pdf 文本解析。")
    add_p(doc, "测试覆盖登录、上传、查重、详情、统计、导出、相似度计算、风险等级、空文本、相似句匹配、UTF-8/GBK 文本解析、Word/PDF 文本解析和不支持类型异常等内容。后端 mvn test 与前端 npm run build 均已通过，说明系统已经能够完成从任务创建、报告上传、文本解析、发起查重到结果展示的完整闭环，满足课程设计的演示要求。")
    heading(doc, "9 总结与展望", 1)
    add_p(doc, "本文完成了一套面向高校实验报告场景的轻量级查重与分析系统。系统不是孤立的算法演示，而是围绕课程、班级、任务和报告建立完整业务模型，并通过 TF-IDF 余弦相似度与 SimHash 组合方法生成可解释的查重结果。")
    add_p(doc, "后续可以继续从中文分词精度、语义相似度、密码加密存储、教师细粒度数据权限、批量上传逐文件反馈、扫描 PDF 的 OCR 处理、前端页面拆分、Word/PDF 报告导出和操作日志等方面完善系统。总体而言，当前系统已经完成开题报告中提出的核心目标，能够作为课程设计成果进行演示和答辩。")
    add_common_references(doc)
    doc.add_page_break()
    heading(doc, "附    件", 1, center=True)
    heading(doc, "附件 A 项目材料说明", 2)
    add_p(doc, "项目材料由后端工程、前端工程、数据库脚本、样例报告和说明文档组成。backend 目录保存 Spring Boot 后端代码，frontend 目录保存 Vue 前端代码，database/init.sql 保存建表语句和初始化演示数据，database/demo_seed.sql 可选预置查重结果和相似句，samples/generated 保存 txt、docx、pdf 三类批量上传样例。")
    add_p(doc, "初始化数据包含 1 名管理员、3 名教师、3 个班级、12 名学生、3 门课程和 3 个主演示实验任务。每个主演示任务均准备 12 份可打开报告样例，并覆盖高风险、中风险、低风险和正常四类对比场景。")
    doc.save(FINAL_OUT)
    return FINAL_OUT


if __name__ == "__main__":
    print(build_initial())
    print(build_final())
