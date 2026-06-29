from pathlib import Path
import math
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "实验报告智能查重与分析系统_课程设计报告终稿.docx"
ASSET_DIR = ROOT / "docs" / "report_assets"
ASSET_DIR.mkdir(exist_ok=True)
SCREENSHOT_DIR = ROOT.parent / "综合课程设计系统截图"

TITLE = "基于 Spring Boot 与文本相似度算法的实验报告查重系统设计与实现"


def font_path(name):
    for p in [Path("C:/Windows/Fonts") / name, Path("C:/Windows/Fonts") / name.lower()]:
        if p.exists():
            return str(p)
    return None


CN_FONT = font_path("simsun.ttc") or font_path("msyh.ttc") or font_path("arial.ttf")


def img_font(size):
    try:
        return ImageFont.truetype(CN_FONT, size)
    except Exception:
        return ImageFont.load_default()


def draw_wrapped(draw, text, box, font, fill=(0, 0, 0), spacing=6):
    x1, y1, x2, y2 = box
    max_w = x2 - x1 - 18
    lines = []
    for para in text.split("\n"):
        line = ""
        for ch in para:
            trial = line + ch
            if draw.textbbox((0, 0), trial, font=font)[2] <= max_w:
                line = trial
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
    heights = [draw.textbbox((0, 0), line, font=font)[3] - draw.textbbox((0, 0), line, font=font)[1] for line in lines]
    total_h = sum(heights) + max(0, len(lines) - 1) * spacing
    y = y1 + max(0, (y2 - y1 - total_h) / 2)
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += h + spacing


def rounded_box(draw, box, text, fill, font, outline=(0, 0, 0)):
    draw.rounded_rectangle(box, radius=14, fill="white", outline=(0, 0, 0), width=2)
    draw_wrapped(draw, text, box, font)


def arrow(draw, start, end, fill=(0, 0, 0), width=4):
    draw.line([start, end], fill=fill, width=width)
    sx, sy = start
    ex, ey = end
    ang = math.atan2(ey - sy, ex - sx)
    size = 14
    pts = [
        (ex, ey),
        (ex - size * math.cos(ang - math.pi / 6), ey - size * math.sin(ang - math.pi / 6)),
        (ex - size * math.cos(ang + math.pi / 6), ey - size * math.sin(ang + math.pi / 6)),
    ]
    draw.polygon(pts, fill=fill)


def make_architecture(path):
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    titlef, boxf, small = img_font(40), img_font(27), img_font(22)
    d.text((650, 40), "系统总体架构", font=titlef, fill=(0, 0, 0))
    boxes = [
        ((70, 180, 330, 330), "用户浏览器\nVue 3 + Element Plus", "#E8F1FB"),
        ((500, 180, 760, 330), "前端工程\nAxios / Router / ECharts", "#EEF7ED"),
        ((930, 180, 1230, 330), "后端服务\nSpring Boot REST API", "#F4F6F9"),
        ((930, 500, 1230, 650), "业务层\n认证/课程/任务/报告/查重", "#F8F0E6"),
        ((70, 500, 330, 650), "文件存储\nuploads 目录", "#F7EBEF"),
        ((500, 500, 760, 650), "MySQL 数据库\n业务数据存储", "#EEF2FF"),
        ((1280, 500, 1530, 650), "算法模块\nTF-IDF / SimHash", "#F2F8F5"),
    ]
    for box, text, color in boxes:
        rounded_box(d, box, text, color, boxf)
    arrow(d, (330, 255), (500, 255))
    arrow(d, (760, 255), (930, 255))
    arrow(d, (1080, 330), (1080, 500))
    arrow(d, (930, 625), (760, 625))
    arrow(d, (930, 630), (330, 630))
    arrow(d, (1230, 575), (1280, 575))
    d.text((455, 360), "HTTP/JSON 接口", font=small, fill=(0, 0, 0))
    d.text((825, 360), "JWT 认证", font=small, fill=(0, 0, 0))
    img.save(path)


def make_function(path):
    img = Image.new("RGB", (1600, 980), "white")
    d = ImageDraw.Draw(img)
    titlef, boxf, small = img_font(40), img_font(25), img_font(21)
    d.text((650, 35), "系统功能结构", font=titlef, fill=(0, 0, 0))
    rounded_box(d, (610, 110, 990, 185), "实验报告智能查重与分析系统", "#E8F1FB", boxf)
    cols = [
        (120, "管理员端", ["用户管理", "课程管理", "班级管理", "班级学生分配", "系统看板"]),
        (520, "教师端", ["实验任务管理", "报告上传与批量导入", "发起查重任务", "查看查重结果", "统计分析与导出"]),
        (920, "学生端", ["查看实验任务", "提交实验报告", "查看提交状态"]),
        (1220, "核心算法与数据", ["文本解析", "文本预处理", "TF-IDF 余弦相似度", "SimHash 指纹", "风险等级与相似句"]),
    ]
    for x, title, items in cols:
        rounded_box(d, (x, 280, x + 260, 350), title, "#F2F4F7", boxf)
        arrow(d, (800, 185), (x + 130, 280))
        y = 400
        for item in items:
            rounded_box(d, (x, y, x + 260, y + 58), item, "#FFFFFF", small, outline=(120, 145, 170))
            y += 78
    img.save(path)


def make_algorithm(path):
    img = Image.new("RGB", (1700, 880), "white")
    d = ImageDraw.Draw(img)
    titlef, boxf, formula = img_font(40), img_font(24), img_font(27)
    d.text((735, 35), "查重算法流程", font=titlef, fill=(0, 0, 0))
    steps = ["读取同一实验任务\n已解析报告", "文本清洗\n去除符号与空白", "二元词条切分\n停用词过滤", "基于报告集合\n计算 TF-IDF 权重", "两两计算\n余弦相似度", "生成 SimHash\n计算指纹相似度", "加权融合\n综合相似度", "风险等级判定\n相似句保存"]
    x, y, w, h = 70, 250, 180, 120
    for i, step in enumerate(steps):
        box = (x + i * (w + 28), y, x + i * (w + 28) + w, y + h)
        rounded_box(d, box, step, "#EAF3F9" if i % 2 == 0 else "#F8F2EA", boxf)
        if i < len(steps) - 1:
            arrow(d, (box[2], y + h / 2), (box[2] + 28, y + h / 2))
    d.rounded_rectangle((270, 520, 1430, 700), radius=16, fill="white", outline=(0, 0, 0), width=2)
    draw_wrapped(d, "综合相似度 = 0.7 × 余弦相似度 + 0.3 × SimHash 相似度\n风险等级：<40% 正常，40%-60% 低风险，60%-80% 中风险，>=80% 高风险", (300, 540, 1400, 680), formula)
    img.save(path)


def make_er(path):
    img = Image.new("RGB", (1700, 1050), "white")
    d = ImageDraw.Draw(img)
    titlef, boxf, small = img_font(40), img_font(23), img_font(18)
    d.text((590, 35), "数据库核心 ER 关系简图", font=titlef, fill=(0, 0, 0))
    nodes = {
        "角色": (110, 160, 320, 250),
        "用户": (430, 160, 640, 250),
        "课程": (750, 160, 960, 250),
        "班级": (1070, 160, 1280, 250),
        "学生班级关联": (1370, 160, 1580, 250),
        "实验任务": (750, 410, 960, 500),
        "报告文件": (1070, 410, 1280, 500),
        "查重任务": (750, 660, 960, 750),
        "查重结果": (1070, 660, 1280, 750),
        "相似句子": (1070, 870, 1280, 960),
    }
    for text, box in nodes.items():
        rounded_box(d, box, text, "white", boxf)

    def mid_label(text, xy):
        bbox = d.textbbox((0, 0), text, font=small)
        x, y = xy
        pad = 5
        d.rectangle((x - pad, y - pad, x + bbox[2] + pad, y + bbox[3] + pad), fill="white")
        d.text((x, y), text, font=small, fill=(0, 0, 0))

    def line(points, label=None, label_xy=None):
        d.line(points, fill=(0, 0, 0), width=3)
        if label and label_xy:
            mid_label(label, label_xy)

    line([(320, 205), (430, 205)], "1:N", (365, 178))
    line([(640, 205), (750, 205)], "教师 1:N", (655, 178))
    line([(1280, 205), (1370, 205)], "1:N", (1308, 178))
    line([(1175, 250), (1175, 410)], "1:N", (1190, 320))
    line([(855, 250), (855, 410)], "1:N", (870, 320))
    line([(960, 455), (1070, 455)], "1:N", (994, 428))
    line([(855, 500), (855, 660)], "1:N", (870, 575))
    line([(960, 705), (1070, 705)], "1:N", (994, 678))
    line([(1175, 500), (1175, 660)], "源/目标报告 1:N", (1190, 575))
    line([(1175, 750), (1175, 870)], "1:N", (1190, 805))
    line([(535, 250), (535, 345), (1175, 345), (1175, 410)], "学生 1:N", (720, 318))
    line([(535, 250), (535, 305), (1475, 305), (1475, 250)], "N:M 通过关联表", (920, 278))
    d.text((100, 985), "说明：用户表通过角色区分管理员、教师和学生；学生与班级通过关联表形成多对多关系；报告文件、查重任务、查重结果和相似句子共同支撑结果追溯。", font=small, fill=(0, 0, 0))
    img.save(path)


def make_business(path):
    img = Image.new("RGB", (1650, 900), "white")
    d = ImageDraw.Draw(img)
    titlef, boxf, note = img_font(40), img_font(25), img_font(25)
    d.text((620, 35), "教师端查重业务流程", font=titlef, fill=(0, 0, 0))
    steps = ["登录系统", "创建课程/班级/任务", "上传或批量导入报告", "系统解析文本\n统计字数与句子数", "发起查重任务", "生成相似度\n风险等级与相似句", "查看结果与图表", "导出 CSV\n人工复核"]
    x, y, w, h = 80, 250, 170, 115
    for i, step in enumerate(steps):
        box = (x + i * (w + 25), y, x + i * (w + 25) + w, y + h)
        rounded_box(d, box, step, "#EEF7ED" if i % 2 == 0 else "#F7EBEF", boxf)
        if i < len(steps) - 1:
            arrow(d, (box[2], y + h / 2), (box[2] + 25, y + h / 2))
    d.rounded_rectangle((220, 560, 1430, 700), radius=16, fill="white", outline=(0, 0, 0), width=2)
    draw_wrapped(d, "系统定位为辅助复核工具：只提供相似度、风险等级和相似句依据，不直接判定学生是否抄袭。", (250, 585, 1400, 680), note)
    img.save(path)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(*(color or (0, 0, 0)))


def set_para(p, line=1.5, before=0, after=0, align=None, first=True):
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Cm(0.74) if first else Cm(0)
    if align is not None:
        p.alignment = align


def add_p(doc, text="", first=True, after=0):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        set_run_font(r, 12)
    set_para(p, first=first, after=after)
    return p


def title_p(doc, text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size, bold=True)
    set_para(p, after=10, first=False)
    return p


def center_heading1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 16, bold=True)
    set_para(p, after=10, first=False)
    return p


def heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, {1: 16, 2: 14, 3: 12}.get(level, 12), bold=True)
    set_para(p, before=8 if level == 1 else 4, after=4, first=False)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, 10.5)
    set_para(p, line=1.2, after=6, first=False)


def set_cell(cell, text, bold=False, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align or (WD_ALIGN_PARAGRAPH.LEFT if len(str(text)) > 14 else WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(str(text))
    set_run_font(r, 10.5, bold=bold)
    set_para(p, line=1.15, first=False)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(table):
    tbl_pr = table._tbl.tblPr
    b = tbl_pr.first_child_found_in("w:tblBorders")
    if b is None:
        b = OxmlElement("w:tblBorders")
        tbl_pr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = b.find(qn("w:" + edge))
        if e is None:
            e = OxmlElement("w:" + edge)
            b.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:color"), "888888")


def add_table(doc, headers, rows, cap):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    borders(table)
    caption(doc, cap)
    return table


def add_pic(doc, path, cap):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(14.5))
    set_para(p, line=1.0, after=2, first=False)
    caption(doc, cap)


def screenshot(name):
    path = SCREENSHOT_DIR / name
    return path if path.exists() else None


def add_screenshot(doc, name, cap):
    path = screenshot(name)
    if path:
        add_pic(doc, path, cap)


def placeholder(doc, cap):
    name = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "_", cap)[:36]
    path = ASSET_DIR / f"placeholder_{name}.png"
    img = Image.new("RGB", (1600, 640), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((30, 30, 1570, 610), outline=(0, 0, 0), width=4)
    d.rectangle((100, 90, 1500, 250), outline=(0, 0, 0), width=3)
    draw_wrapped(d, "系统截图占位\n请后续替换为实际运行界面截图", (100, 105, 1500, 235), img_font(34), fill=(0, 0, 0))
    draw_wrapped(d, cap, (120, 340, 1480, 440), img_font(25), fill=(0, 0, 0))
    img.save(path)
    add_pic(doc, path, cap)


def section_setup(sec):
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.75)
    sec._sectPr.pgMar.set(qn("w:gutter"), str(int(Cm(1.0).twips)))


def add_header_footer(sec, roman=False):
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    for p in sec.header.paragraphs:
        p.clear()
    for p in sec.footer.paragraphs:
        p.clear()
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run(TITLE)
    set_run_font(r, 10.5)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    rr = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "I" if roman else "1"
    rr.append(t)
    fld.append(rr)
    fp._p.append(fld)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), "upperRoman" if roman else "decimal")
    pg.set(qn("w:start"), "1")
    sec._sectPr.append(pg)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for el in [
        ("w:fldChar", {"w:fldCharType": "begin"}),
        ("w:instrText", {"xml:space": "preserve"}, 'TOC \\o "1-3" \\h \\z \\u'),
        ("w:fldChar", {"w:fldCharType": "separate"}),
        ("w:fldChar", {"w:fldCharType": "end"}),
    ]:
        node = OxmlElement(el[0])
        for k, v in el[1].items():
            node.set(qn(k), v)
        if len(el) > 2:
            node.text = el[2]
        run._r.append(node)
    set_para(p, first=False)


def build_doc():
    figs = {
        "architecture": ASSET_DIR / "fig_architecture.png",
        "function": ASSET_DIR / "fig_function.png",
        "algorithm": ASSET_DIR / "fig_algorithm.png",
        "er": ASSET_DIR / "fig_er.png",
        "business": ASSET_DIR / "fig_business.png",
    }
    make_architecture(figs["architecture"])
    make_function(figs["function"])
    make_algorithm(figs["algorithm"])
    make_er(figs["er"])
    make_business(figs["business"])

    doc = Document()
    section_setup(doc.sections[0])
    doc.sections[0].header.is_linked_to_previous = False
    doc.sections[0].footer.is_linked_to_previous = False
    for p in doc.sections[0].header.paragraphs:
        p.clear()
    for p in doc.sections[0].footer.paragraphs:
        p.clear()
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[style_name]
        st.font.name = "Times New Roman"
        st.font.color.rgb = RGBColor(0, 0, 0)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    doc.styles["Normal"].font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("评阅教师：____________")
    set_run_font(r, 12)
    set_para(p, first=False)
    doc.add_paragraph()
    title_p(doc, "海南大学计算机科学与技术学院", 16)
    title_p(doc, "计算机综合课程设计报告", 22)
    for _ in range(3):
        doc.add_paragraph()
    title_p(doc, TITLE, 18)
    for label, value in [
        ("班    级", "计算机科学与技术 2023级"),
        ("姓    名", "（请填写）"),
        ("学    号", "（请填写）"),
        ("组    员", "无"),
        ("指导老师", "（请填写）"),
        ("完成日期", "2026年6月"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}：        {value}        ")
        set_run_font(r, 14)
        set_para(p, line=1.8, after=4, first=False)

    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_setup(sec)
    add_header_footer(sec, roman=True)
    title_p(doc, "摘    要", 16)
    add_p(doc, "随着高校实验教学和课程实践规模不断扩大，实验报告成为评价学生实践能力和课程学习效果的重要依据。传统人工检查方式在面对同一实验任务下的大批量报告时效率较低，且难以及时发现大段复制、模板套用和局部改写等问题。针对这一场景，本文设计并实现了一套基于 Spring Boot 与文本相似度算法的实验报告查重系统。系统采用前后端分离架构，后端基于 Spring Boot、MyBatis-Plus 和 MySQL 实现用户认证、课程管理、班级管理、实验任务管理、报告上传、文本解析、查重计算和结果导出等功能，前端基于 Vue 3、Element Plus 和 ECharts 实现后台管理界面、统计看板和查重结果展示。在算法方面，系统支持 txt、docx 和 pdf 报告文本提取，对文本进行清洗、轻量级中文二元词条切分和停用词过滤，并以同一实验任务下的全部已解析报告作为语料集合计算 TF-IDF 权重。系统通过余弦相似度衡量报告向量之间的相似程度，同时引入 SimHash 文本指纹方法计算近重复文本相似度，最终按照加权融合方式得到综合相似度，并划分正常、低风险、中风险和高风险等级。系统还提供相似句子展示、统计图表、CSV 结果导出和多班级多任务演示数据，能够为教师批量复核实验报告提供可量化、可解释的辅助依据。测试结果表明，系统主流程完整，能够完成从任务创建、报告上传、文本解析、发起查重到结果展示的闭环操作，满足课程设计预期目标。")
    p = doc.add_paragraph()
    r = p.add_run("关键词：")
    set_run_font(r, 12, bold=True)
    r = p.add_run("实验报告查重；文本相似度；TF-IDF；余弦相似度；SimHash")
    set_run_font(r, 12)
    set_para(p, first=False)

    doc.add_page_break()
    title_p(doc, "Abstract", 16)
    add_p(doc, "With the expansion of experimental teaching and practical courses in universities, experimental reports have become important materials for evaluating students' practical ability and learning outcomes. Manual review is inefficient when teachers need to inspect a large number of reports under the same experimental task, and it is difficult to identify large-scale copying, template reuse and partial rewriting in time. To address this problem, this paper designs and implements an experimental report plagiarism checking system based on Spring Boot and text similarity algorithms. The system adopts a front-end and back-end separated architecture. The back end is built with Spring Boot, MyBatis-Plus and MySQL, implementing user authentication, course management, class management, experiment task management, report upload, text extraction, similarity calculation and result export. The front end is built with Vue 3, Element Plus and ECharts, providing management pages, dashboards and result visualization. In terms of algorithms, the system extracts text from txt, docx and pdf files, performs text cleaning, lightweight Chinese bigram tokenization and stop-word filtering, and calculates TF-IDF weights based on all parsed reports under the same task. Cosine similarity is used to measure vector similarity, while SimHash is introduced to evaluate near-duplicate text fingerprints. The final similarity score is obtained through weighted fusion and mapped to normal, low-risk, medium-risk and high-risk levels. The system also provides similar sentence display, statistical charts, CSV export and multi-class demonstration data, offering quantitative and interpretable support for teachers. Test results show that the system can complete the full workflow from task creation, report upload and text parsing to plagiarism checking and result display, meeting the expected objectives of the course project.")
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    set_run_font(r, 12, bold=True)
    r = p.add_run("Experimental report checking; Text similarity; TF-IDF; Cosine similarity; SimHash")
    set_run_font(r, 12)
    set_para(p, first=False)

    doc.add_page_break()
    title_p(doc, "目    录", 16)
    add_toc(doc)

    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section_setup(sec)
    add_header_footer(sec, roman=False)

    chapters = [
        ("1 绪论", [
            ("1.1 研究背景及意义", [
                "随着高校实验教学、课程设计和综合实践教学规模不断扩大，实验报告已经成为记录学生实践过程、反映课程理解程度和评价动手能力的重要材料。以计算机专业课程为例，一门课程往往包含多个实验任务，每个班级会集中提交数十份甚至上百份报告。教师在批阅报告时，不仅需要关注实验过程、结果分析和格式规范，还需要判断报告内容是否存在大段复制、简单改写、模板套用或网络资料拼接等情况。传统人工检查方式依赖教师经验，面对批量报告时效率较低，难以及时定位疑似重复内容。",
                "通用论文查重系统通常面向毕业论文、期刊论文等长文本，使用成本较高，且不一定适合实验报告这种篇幅较短、结构相对固定、按课程任务提交的教学场景。因此，设计一套面向高校实验教学的轻量级实验报告查重系统具有实际意义。该系统不直接判定学生是否抄袭，而是通过相似度、风险等级和相似句子等信息，为教师提供可量化、可复核的辅助依据。",
                "从课程设计角度看，本课题综合运用了 Web 后端开发、前端交互设计、数据库建模、文件解析、文本预处理和相似度算法等知识，能够体现计算机专业课程综合实践的特点。系统围绕课程、班级、实验任务、报告文件和查重结果形成完整业务链路，相比单独实现一个文本相似度计算程序，更符合实际教学管理需求。"
            ]),
            ("1.2 国内外研究现状", [
                "文本相似度计算是自然语言处理、信息检索和文档管理领域的重要研究内容。传统方法主要包括词袋模型、n-gram、Jaccard 相似度、TF-IDF、余弦相似度和局部敏感哈希等。这类方法实现成本较低、计算效率较高、结果解释性较强，适合中小规模教学系统中的近重复文本检测。",
                "近年来，文本相似度研究逐渐从统计特征扩展到语义特征和深度学习模型。基于 Transformer、句向量和预训练语言模型的方法在语义匹配和改写识别方面具有优势，但模型部署、算力消耗和数据依赖较高。结合课程设计周期和系统部署成本，本系统选择 TF-IDF、余弦相似度和 SimHash 作为主要算法，兼顾可实现性、运行效率和可解释性。",
                "对于中文实验报告而言，文本处理还会受到中文分词、专业术语、固定实验模板、代码片段和实验步骤相似等因素影响。为降低实现复杂度，系统采用轻量级中文二元词条切分方法，并结合停用词过滤完成文本预处理。该方法不能替代成熟语义模型，但能够在课程设计范围内满足实验报告近重复检测需求。"
            ]),
            ("1.3 本文主要研究内容", [
                "本文围绕“基于 Spring Boot 与文本相似度算法的实验报告查重系统设计与实现”展开，主要完成以下工作：第一，分析高校实验报告查重场景的业务需求，确定管理员、教师和学生三类角色及其功能边界；第二，设计基于 Spring Boot、Vue 和 MySQL 的前后端分离系统架构；第三，建立课程、班级、实验任务、报告文件、查重任务、查重结果和相似句子等核心数据表；第四，实现 txt、docx、pdf 文本解析和上传内容基础检测；第五，设计并实现基于 TF-IDF 余弦相似度与 SimHash 的组合查重算法；第六，完成查重结果列表、相似句详情、统计图表和 CSV 导出等功能；第七，设计多班级、多学生、多任务演示数据并进行系统测试。"
            ]),
            ("1.4 论文组织结构", [
                "全文共分为九章。第 1 章介绍研究背景、意义、研究现状和本文主要工作；第 2 章介绍系统相关理论基础与关键技术；第 3 章进行需求分析；第 4 章给出系统总体设计；第 5 章说明数据库设计；第 6 章介绍查重算法设计与实现；第 7 章介绍系统主要功能实现；第 8 章进行系统测试与结果分析；第 9 章总结全文并提出后续展望。"
            ])
        ]),
        ("2 理论基础与关键技术", [
            ("2.1 B/S 架构与前后端分离", [
                "B/S 架构即 Browser/Server 架构，用户通过浏览器访问系统，主要业务逻辑部署在服务器端。该架构不需要在用户端安装专门客户端，便于系统部署、升级和维护。实验报告查重系统面向教师和学生使用，采用 B/S 架构能够降低使用门槛，用户只需通过浏览器即可完成登录、任务查看、报告上传和结果查看等操作。",
                "本系统采用前后端分离模式，前端负责页面展示、交互控制和图表渲染，后端提供 REST API、业务逻辑和数据持久化。前端使用 Axios 调用后端接口，后端统一返回 JSON 数据。前后端分离使系统结构更加清晰，也便于后续扩展移动端或其他客户端。"
            ]),
            ("2.2 Spring Boot 与 MyBatis-Plus", [
                "Spring Boot 是基于 Spring 生态的快速开发框架，能够通过自动配置简化 Web 应用开发。本系统后端使用 Spring Boot 构建 REST API，并通过 Controller、Service、Mapper 分层组织代码。Controller 层负责接收请求，Service 层处理业务逻辑，Mapper 层负责数据库访问。",
                "MyBatis-Plus 是 MyBatis 的增强工具，提供常用 CRUD、条件构造器和分页等能力。本系统使用 MyBatis-Plus 操作 MySQL 数据库，可以减少重复 SQL 编写，提高开发效率。系统中的用户、课程、班级、实验任务、报告文件和查重结果均通过实体类和 Mapper 完成持久化。"
            ]),
            ("2.3 Vue、Element Plus 与 ECharts", [
                "Vue 3 是渐进式 JavaScript 框架，适合构建响应式单页应用。Element Plus 提供表格、表单、对话框、菜单等后台管理常用组件，可以提高前端开发效率。ECharts 是常用的数据可视化库，本系统使用它展示风险等级占比、核心数据统计、相似度区间分布和查重任务趋势等图表。",
                "系统前端采用单页后台管理结构，通过侧边栏菜单组织首页看板、用户管理、课程管理、班级管理、实验任务、报告上传、查重结果和统计分析等功能。虽然当前主要功能集中在一个视图文件中，但从用户使用角度已经形成清晰的模块划分。"
            ]),
            ("2.4 文本相似度算法基础", [
                "TF-IDF 是信息检索中常用的词项权重计算方法，用于衡量词项在文档集合中的重要程度。词频 TF 反映词项在当前文档中的出现频率，逆文档频率 IDF 反映词项在整个文档集合中的区分能力。对于实验报告查重场景，若某些词语只在少数报告中频繁出现，则这些词语对区分文本更有价值。",
                "余弦相似度通过计算两个向量夹角的余弦值衡量相似程度，取值越接近 1 表示两个文本向量越接近。SimHash 是一种文本指纹算法，可以将文本特征映射为固定长度的二进制指纹，再通过汉明距离判断文本指纹的接近程度。余弦相似度适合衡量词项重合程度，SimHash 适合快速判断近重复文本，二者结合能够提高结果的稳定性和解释性。"
            ])
        ]),
        ("3 需求分析", [
            ("3.1 用户角色分析", [
                "系统主要面向管理员、教师和学生三类用户。管理员负责基础数据维护；教师负责实验任务、报告上传、查重任务和结果分析；学生负责查看实验任务并提交自己的实验报告。不同角色具有不同功能边界，系统通过 JWT 和角色字段进行识别。"
            ]),
            ("3.2 功能需求分析", [
                "系统功能需求可以分为登录认证、基础数据管理、实验任务管理、报告上传与解析、查重计算、结果展示、统计分析和结果导出等模块。登录认证模块负责用户身份验证和 Token 生成；基础数据管理模块维护用户、课程、班级和学生班级关系；实验任务管理模块用于创建和维护课程实验任务；报告上传模块支持单文件上传和批量上传；查重模块对同一任务下报告进行两两比较；结果模块展示相似度、风险等级和相似句；统计模块展示整体数据和风险分布。",
                "报告上传还需要完成基础内容检测，包括文件是否为空、文件大小是否超过限制、文件类型是否合法、文本是否能够提取、字数统计、句子数量统计以及文本过短提示。该检测结果显示在报告列表中，便于教师了解上传文件质量。"
            ]),
            ("3.3 非功能需求分析", [
                "在安全性方面，系统需要实现基础登录认证、角色权限控制、上传文件类型校验和文件大小限制；在可维护性方面，后端应采用分层结构，前端应按功能模块组织页面逻辑；在可扩展性方面，系统应支持后续增加更多文件格式、分词算法、历史报告库和 Word/PDF 报告导出；在易用性方面，系统应提供清晰的后台菜单、直观的风险标签和统计图表。"
            ])
        ]),
    ]

    chapter_intros = {
        "1 绪论": "本章主要说明课题产生的背景、研究意义、相关研究基础以及本文的主要工作。通过对高校实验报告批量复核场景的分析，可以明确本系统并不是通用论文查重平台，而是面向课程实验任务的轻量级辅助查重工具。",
        "2 理论基础与关键技术": "本章围绕系统实现所需的技术基础展开，先介绍 Web 系统采用的 B/S 架构和前后端分离模式，再说明后端、前端、数据库和文本相似度算法的选择依据，为后续系统设计和实现提供技术支撑。",
        "3 需求分析": "本章从实际使用角色和业务流程出发分析系统需求。系统需要同时满足管理员维护基础数据、教师组织实验与复核报告、学生提交实验报告三类场景，因此需求分析不仅关注页面功能，还需要关注权限边界、上传约束和结果解释方式。"
    }

    for chap, subsections in chapters:
        heading(doc, chap, 1)
        add_p(doc, chapter_intros[chap])
        for sub, paras in subsections:
            heading(doc, sub, 2)
            for para in paras:
                add_p(doc, para)
        if chap.startswith("3 "):
            add_p(doc, "综合上述需求，系统没有把查重算法孤立为单一工具，而是将其放入课程、班级、任务和报告管理的完整业务链路中。管理员主要负责用户、课程、班级和班级学生分配等基础数据；教师负责实验任务、报告上传、批量导入、查重发起、结果复核、统计分析和 CSV 导出；学生则完成任务查看和本人报告提交。这样的角色划分既能支撑答辩演示，也便于后续按权限控制接口访问范围。")
            add_p(doc, "从开题报告承诺到代码实现，系统已经覆盖登录认证、基础数据管理、实验任务管理、报告上传解析、查重计算、结果展示、统计分析和结果导出等核心模块。后续章节的总体设计、数据库设计和功能实现均围绕这些模块展开。")

    heading(doc, "4 系统总体设计", 1)
    add_p(doc, "本章在需求分析基础上给出系统总体设计。系统采用前后端分离结构，将浏览器交互、后端业务处理、数据库持久化、文件存储和查重算法组合成完整业务链路。为了避免设计说明停留在文字层面，本章通过系统架构图、功能结构图和教师端业务流程图说明各模块之间的协作关系。")
    heading(doc, "4.1 系统架构设计", 2)
    add_p(doc, "系统采用前后端分离架构。前端基于 Vue 3 构建后台管理页面，通过 Axios 调用后端接口；后端基于 Spring Boot 提供 REST API，并通过 MyBatis-Plus 访问 MySQL 数据库；报告文件保存到本地 uploads 目录；查重算法模块在后端服务中完成文本预处理、相似度计算和结果入库。为了避免系统结构说明停留在组件罗列层面，本文将浏览器访问、前端请求、后端业务、数据库持久化、文件存储和算法处理之间的调用关系整理为图4-1。")
    add_pic(doc, figs["architecture"], "图4-1 系统总体架构图")
    add_p(doc, "图4-1 按访问展示、后端业务、数据与算法三个层次展开系统结构。可以看到，前端只负责页面交互和接口调用，不直接操作数据库或上传目录；后端既承担身份认证、课程班级、实验任务和报告管理等业务逻辑，也负责在查重任务触发后统一读取报告文本并写入结果。这种组织方式使权限校验、文件处理和算法计算都集中在后端服务中，能够减少前端绕过业务规则访问数据的风险，也便于后续扩展新的报告格式或查重策略。")
    heading(doc, "4.2 功能结构设计", 2)
    add_p(doc, "在明确总体架构之后，还需要进一步说明系统功能如何落到不同角色上。管理员侧重基础数据维护，教师侧重实验任务与查重分析，学生侧重任务查看和报告提交，核心算法和数据模块则为这些业务功能提供支撑。图4-2 将功能按管理员端、教师端、学生端和算法数据支撑四个部分展开，用于检查开题报告中承诺的功能是否在系统设计层面形成闭环。")
    add_pic(doc, figs["function"], "图4-2 系统功能结构图")
    add_p(doc, "由图4-2 可以看出，系统并不是单纯提供一个相似度计算入口，而是把查重功能嵌入课程、班级、实验任务和报告提交的教学业务中。管理员维护基础数据后，教师才能创建任务并组织报告上传；学生提交报告后，算法模块才能基于同一任务范围生成可复核结果。因此，教师端的任务创建、批量上传、发起查重、结果查看和导出构成答辩演示主线，管理员端和学生端则共同保证演示数据来源完整。")
    heading(doc, "4.3 业务流程设计", 2)
    add_p(doc, "教师端业务流程是系统的核心流程。教师登录后创建实验任务，学生或教师上传报告，系统解析报告文本并记录字数和句子数。教师发起查重任务后，系统对同一任务下所有解析成功报告进行两两比较，保存查重结果和相似句子，最后在前端展示风险等级和统计图表。为了说明这些操作不是彼此割裂的页面功能，图4-3 按教师实际使用顺序展示了从任务准备到结果复核的完整链路。")
    add_pic(doc, figs["business"], "图4-3 教师端查重业务流程图")
    add_p(doc, "从图4-3 的顺序可以看出，报告解析和查重计算处于流程中段，前面依赖课程、班级和任务数据，后面服务于结果列表、相似句详情、统计图表和 CSV 导出。也就是说，系统输出的风险等级并不是自动给出最终判定，而是把教师原本需要逐份比较的工作转化为可排序、可追溯、可复核的结果集合，最终仍由教师结合任务模板和报告内容进行判断。")
    heading(doc, "4.4 项目目录结构设计", 2)
    add_p(doc, "在工程实现层面，最终提交源码按照后端、前端、数据库脚本和样例数据进行组织。其中 backend 是 Spring Boot 后端项目，包含控制器、服务、实体、Mapper、算法和工具类；frontend 是 Vue 3 前端项目，包含页面视图、接口封装、路由和样式；database 保存统一的 MySQL 建表和演示数据脚本；samples 保存用于批量上传和查重演示的样例报告。课程设计报告和过程文档单独放在文档提交目录中，不混入源码压缩包。这样的目录划分可以使运行代码、数据库脚本、演示数据和论文材料相互独立，降低评阅和复现成本。")

    heading(doc, "5 数据库设计", 1)
    add_p(doc, "本章说明系统的数据组织方式。实验报告查重业务需要同时保存用户身份、课程班级、实验任务、上传报告、查重结果和相似句等信息，因此数据库设计的重点是让查重结果能够追溯到具体任务、学生和报告文件。正文中先给出简化实体关系图，再对核心数据表进行中文说明。")
    heading(doc, "5.1 数据库设计原则", 2)
    add_p(doc, "数据库设计围绕实验报告查重业务展开，遵循实体清晰、关系明确和便于扩展的原则。系统使用 MySQL 作为数据存储，主要保存用户、角色、课程、班级、实验任务、报告文件、查重任务、查重结果和相似句子等信息。")
    add_p(doc, "其中，report_file 表保存上传报告的文件信息、解析文本、字数和解析状态；check_task 表保存一次查重任务的执行状态和报告数量；check_result 表保存两份报告之间的余弦相似度、SimHash 相似度、综合相似度和风险等级；similar_sentence 表保存结果详情中的相似句子。")
    add_p(doc, "由于查重结果需要回溯到具体任务、学生和报告文件，仅用文字描述表关系不够直观，因此本文将核心实体之间的关系整理为图5-1。该图与真实数据库表保持一致，包含 sys_role、sys_user、course、class_info、student_class、experiment_task、report_file、check_task、check_result 和 similar_sentence 等关键表，并在关系线上标注一对多、多对多等主要基数关系。")
    add_pic(doc, figs["er"], "图5-1 数据库核心实体关系图")
    add_p(doc, "图5-1 中，角色与用户是一对多关系，用户与班级在业务上可表示为多对多关系，课程、班级与实验任务共同限定报告提交范围，报告文件再通过比较关系生成多条查重结果。这样的关系设计使系统能够从任意一条查重结果继续追溯到源报告、对比报告、实验任务和提交学生，满足教师复核时对结果来源的追踪需求，也避免查重结果脱离教学任务单独存在。")
    heading(doc, "5.2 核心数据表设计", 2)
    add_p(doc, "ER 图解决的是实体之间的关系问题，而实际落库时还需要说明每类表承担的业务职责。基于图5-1 的核心关系，表5-1 用中文归纳主要数据表保存的内容和作用。该表不是完整建表语句，而是面向论文说明的数据字典摘要，完整字段、主键、外键和初始化数据以 database/report_check_system_full.sql 为准。")
    add_table(doc, ["数据表中文名称", "主要保存内容", "业务作用"], [
        ["用户与角色表", "账号、姓名、角色、学号、工号、状态", "区分管理员、教师和学生，为登录认证和权限控制提供依据"],
        ["课程与班级表", "课程名称、课程编码、任课教师、班级名称、年级、专业", "描述实验教学的基础组织信息"],
        ["学生班级关联表", "学生、班级之间的对应关系", "支持一个班级包含多个学生，也便于后续扩展学生跨班级场景"],
        ["实验任务表", "所属课程、所属班级、任务标题、任务说明、截止时间、状态", "限定报告提交和查重比较的业务范围"],
        ["报告文件表", "提交学生、原始文件名、文件类型、解析文本、字数、解析状态", "保存上传文件及其可用于查重的文本内容"],
        ["查重任务表", "实验任务、算法名称、报告数量、执行状态、开始和结束时间", "记录一次批量查重过程"],
        ["查重结果与相似句表", "源报告、对比报告、综合相似度、风险等级、相似句内容", "保存两份报告的比较结果，并提供教师复核依据"],
    ], "表5-1 核心数据表中文说明")
    add_p(doc, "结合表5-1 可以看出，系统没有把所有查重信息堆放在一个大表中，而是将报告原文、查重任务、相似度结果和相似句子拆分存储。报告文件表负责保存可解析文本，查重任务表记录一次批量计算过程，查重结果与相似句表分别面向列表展示和详情复核。这种拆分既降低了单表字段复杂度，也便于前端按照报告列表、查重结果列表和相似句详情三个层次读取数据。")
    heading(doc, "5.3 演示数据设计", 2)
    add_p(doc, "为保证系统能够真实演示，最终提交的 database/report_check_system_full.sql 将建库建表、基础账号、课程班级、实验任务、报告记录、查重任务、查重结果和相似句样例集中到一个脚本中。执行该脚本后，系统会生成 1 个管理员、3 个教师、12 个学生、3 个班级、3 门课程和 3 个主演示实验任务，并预置可直接展示的查重结果。samples/generated 目录提供 12 份 txt、12 份 docx 和 12 份文字型 pdf 报告，教师端批量上传时可以按文件名自动匹配学生。")

    heading(doc, "6 查重算法设计与实现", 1)
    add_p(doc, "本章说明系统查重算法的处理过程。由于实验报告篇幅相对较短、结构较固定，系统采用可解释性较强的 TF-IDF、余弦相似度和 SimHash 组合方法，而不是直接引入复杂语义模型。算法设计重点是让教师能够理解相似度来源，并通过风险等级和相似句进一步复核。")
    heading(doc, "6.1 文本解析与预处理", 2)
    add_p(doc, "系统支持 txt、docx 和 pdf 三类常见实验报告格式。txt 文件优先按 UTF-8 编码读取，失败时兼容 GBK 编码；docx 文件通过 Apache POI 提取段落文本；pdf 文件通过 PDFBox 提取普通文本内容。解析后的文本保存到 report_file 表中，供后续查重任务使用。")
    add_p(doc, "文本预处理包括文本清洗、大小写统一、去除非中文英文数字字符、轻量级中文二元词条切分和停用词过滤。由于课程设计周期和部署成本限制，系统没有接入复杂中文分词模型，而是采用二元词条切分方法将连续中文文本转换为相邻二字词项。该方法实现简单、计算成本较低，适合实验报告近重复检测。")
    heading(doc, "6.2 TF-IDF 与余弦相似度计算", 2)
    add_p(doc, "系统以同一实验任务下所有解析成功的报告作为语料集合，统计词项在不同报告中的文档频率，并计算 IDF 权重。对于每一份报告，系统根据词频和 IDF 构建文本向量。这样可以避免只在两份文本之间临时计算 IDF，使算法流程更符合开题报告中“基于同一任务报告集合建模”的设计思路。")
    add_p(doc, "TF(t,d)=词项 t 在报告 d 中出现次数 / 报告 d 的词项总数", first=False)
    add_p(doc, "IDF(t)=ln((N+1)/(DF(t)+1))+1", first=False)
    add_p(doc, "Cosine(A,B)=(A·B)/(||A||×||B||)", first=False)
    heading(doc, "6.3 SimHash 文本指纹计算", 2)
    add_p(doc, "SimHash 将文本特征映射为固定长度的指纹。系统对每个词项生成稳定哈希值，再根据词项频率对 64 位向量进行加权累加。向量中每一位大于 0 则置为 1，否则置为 0，最终得到文本指纹。两个文本指纹的汉明距离越小，说明文本整体特征越接近。")
    add_p(doc, "在本系统中，SimHash 相似度计算公式为 1 - 汉明距离 / 64。为避免空文本产生误判，系统对空文本进行了保护处理：当任一文本没有有效词项时，SimHash 相似度直接返回 0。")
    heading(doc, "6.4 综合相似度与风险等级", 2)
    add_p(doc, "系统将余弦相似度和 SimHash 相似度进行加权融合，其中余弦相似度权重为 0.7，SimHash 相似度权重为 0.3。该设计使系统既关注词项向量层面的重合程度，也关注整体文本指纹的近重复程度。")
    add_p(doc, "FinalScore = 0.7 × CosineSimilarity + 0.3 × SimHashSimilarity", first=False)
    add_p(doc, "为了让教师能够快速理解查重结果，系统没有只返回一个孤立的相似度数值，而是将综合相似度映射为四类风险等级。阈值划分的作用不是替代教师判断，而是帮助教师确定复核优先级：相似度越高，越需要优先打开相似句详情进行人工确认。表6-1 给出了系统当前采用的风险等级规则。")
    add_table(doc, ["综合相似度", "风险等级", "说明"], [
        [">= 0.80", "高风险", "文本高度相似，需要教师重点复核"],
        ["0.60 - 0.80", "中风险", "存在明显相似内容，建议查看相似句"],
        ["0.40 - 0.60", "低风险", "存在一定重合，需结合任务模板判断"],
        ["< 0.40", "正常", "未发现明显重复风险"],
    ], "表6-1 风险等级划分规则")
    add_p(doc, "从表6-1 可以看出，低风险和正常结果主要用于减少教师无效检查，高风险和中风险结果则是后续相似句复核的重点。由于实验报告通常包含固定实验目的、步骤模板和代码框架，阈值只能反映文本相似程度，不能直接等同于抄袭结论，因此系统在结果页继续保留相似句详情，让教师能够判断相似内容来自正常模板还是异常重复。")
    add_p(doc, "将文本解析、向量化、指纹计算和风险映射串联起来后，可以形成完整的查重算法流程。图6-1 展示的是一次查重任务在后端内部的处理顺序，重点说明系统如何从同一实验任务下的报告集合出发，逐步得到综合相似度、风险等级和相似句记录。")
    add_pic(doc, figs["algorithm"], "图6-1 查重算法流程图")
    add_p(doc, "图6-1 的流程与代码中的 SimilarityCalculator、ReportFileService 和 CheckTaskService 相对应。报告文本先经过清洗、二元词条切分和停用词过滤，再以同一任务下的报告集合作为语料计算 TF-IDF 权重；随后系统分别计算余弦相似度和 SimHash 相似度，并按照加权融合规则生成综合相似度。相似句保存放在流程末端，是为了让结果不仅有分数，还能提供教师可检查的文本依据。")

    heading(doc, "7 系统功能实现", 1)
    add_p(doc, "本章按照系统演示流程介绍主要功能实现。系统功能从登录认证开始，经过首页统计、基础数据维护、实验任务管理、报告上传与解析、查重任务执行、结果展示和导出，最终形成教师可使用的报告复核闭环。各节内容均对应当前代码中已经实现的业务模块。")
    sections = [
        ("7.1 登录认证与权限控制", ["系统登录模块接收用户名和密码，后端根据 sys_user 表查询用户信息并校验用户状态。登录成功后，JwtUtil 生成包含用户编号、用户名和角色编码的 Token，前端将 Token 保存到本地存储，并在后续请求中通过 Authorization 请求头发送给后端。系统还提供教师和学生注册、邮箱找回密码等入口，管理员账号由后台维护，不允许前台注册。", "后端通过 AuthInterceptor 解析 Token，并将当前用户信息保存到 AuthContext 中。业务层根据角色判断用户权限。例如学生只能查看和提交自己的报告，不能批量上传报告，也不能查看他人的查重结果。登录认证是后续所有业务请求的入口，保证课程、班级、实验任务和查重结果均在角色权限范围内访问。"]),
        ("7.2 首页看板与统计分析", ["首页看板展示课程数量、实验任务数量、报告数量和查重任务数量等核心指标。统计分析页面通过 ECharts 展示风险等级占比、核心数据统计、相似度区间分布和查重任务趋势，帮助教师从整体上了解当前实验任务的报告重复风险。", "首页和统计分析功能不是独立的装饰页面，而是对课程、任务、报告和查重结果数据的汇总。教师可以先通过看板了解整体情况，再进入具体任务查看报告和风险结果。"]),
        ("7.3 用户、课程与班级管理", ["管理员端提供用户管理、课程管理、班级管理和班级学生分配功能。用户管理支持新增、编辑、删除和密码重置；课程管理维护课程名称、课程编码和任课教师；班级管理维护年级、专业和班级名称；班级学生分配用于建立学生和班级之间的关联关系。", "基础数据管理保证了后续查重任务具有明确的课程、班级和学生归属。没有这些数据，查重结果只能停留在文件之间的比较，无法服务于实际教学管理。"]),
        ("7.4 实验任务管理", ["教师可以创建实验任务，设置所属课程、所属班级、任务标题、任务说明和截止时间。系统通过实验任务将报告、查重任务和结果组织起来，保证查重只在同一实验任务下进行，避免不同任务之间因主题差异导致结果失真。", "实验任务是系统组织报告和查重结果的核心单位。系统按任务限定报告集合，可以避免不同实验主题之间进行无意义比较，也便于教师按课程教学进度分批管理报告。"]),
        ("7.5 报告上传与内容检测", ["报告上传模块支持单文件上传和批量上传。教师端可以选择实验任务和学生上传报告，也可以批量选择多个文件。批量上传时，系统会根据文件名中的学生姓名、用户名或学号自动匹配学生。学生端只能提交自己的报告。", "上传后系统会完成文件类型、文件大小、任务状态、学生身份、学生是否属于任务班级、文本提取、字数统计和句子数量统计等检测。若文本过短或句子数量过少，系统会在解析状态中给出“建议教师复核”的提示。"]),
        ("7.6 查重任务与结果展示", ["教师选择实验任务后可以发起查重任务。后端查询该任务下所有解析成功的报告，如果报告数量少于两份，则提示无法查重。符合条件时，系统创建 check_task 记录，并对报告进行两两比较，生成 check_result 和 similar_sentence 数据。", "查重结果列表按照综合相似度降序展示，教师可以查看报告 A、报告 B、余弦相似度、SimHash 相似度、综合相似度和风险等级。点击详情后，系统展示相似句子列表，便于教师复核具体相似内容。", "查重结果列表将算法输出转化为教师能够直接使用的信息，包括相似度数值、风险等级和详情入口。这样设计能够让教师先筛选高风险记录，再查看相似句进行复核。"]),
    ]
    screenshots = {
        "7.1": ("初始登录界面截图.png", "图7-1 系统登录界面"),
        "7.2": ("admin系统看板.png", "图7-2 系统首页看板"),
        "7.3": ("admin用户清单.png", "图7-3 管理员用户管理界面"),
        "7.4": ("teacher01实验任务.png", "图7-4 教师端实验任务管理界面"),
        "7.5": ("teacher01报告上传.png", "图7-5 教师端报告上传与解析结果"),
        "7.6": ("teacher01查重结果.png", "图7-6 查重结果列表界面"),
    }
    for sub, paras in sections:
        heading(doc, sub, 2)
        for para in paras:
            add_p(doc, para)
        for prefix, shot in screenshots.items():
            if sub.startswith(prefix):
                add_screenshot(doc, shot[0], shot[1])
                break
        if sub.startswith("7.5"):
            add_p(doc, "上传模块直接决定后续查重数据是否可靠。如果空文件、格式不支持文件或解析失败文件进入算法流程，最终相似度就会失去解释意义。因此，系统在保存报告记录前后设置了多项基础检测，表7-1 按检测对象、检测内容和处理方式对这些规则进行归纳。")
            add_table(doc, ["检测项", "检测内容", "处理方式"], [
                ["文件为空", "判断 MultipartFile 是否为空", "拒绝上传并提示"],
                ["文件大小", "单个文件不超过 20MB", "超出限制时拒绝上传"],
                ["文件类型", "仅允许 txt、docx、pdf", "不支持类型直接拦截"],
                ["任务状态", "任务需存在且处于进行中", "截止或归档后禁止提交"],
                ["学生身份", "上传对象必须为学生角色", "非学生用户不能作为提交人"],
                ["文本提取", "解析后文本不能为空", "失败时记录解析失败原因"],
                ["内容质量", "统计字数和句子数量", "文本偏短时提示教师复核"],
            ], "表7-1 上传内容检测规则")
            add_p(doc, "表7-1 中的规则覆盖了文件本身、实验任务状态、提交人身份和解析后文本质量四个层面。这样设计的意义在于把问题尽量拦截在上传阶段：明显不合法的文件不进入系统，能够解析但内容偏短的报告则保留记录并提示教师复核，从而兼顾系统自动处理和人工判断。")
    add_p(doc, "相似度分值能够帮助教师排序结果，但无法直接说明两份报告究竟在哪些内容上相似。为补充这一信息，系统在结果详情中提供相似句列表，展示源报告句子、对比报告句子和句子级相似度。教师可以据此定位两份报告中具体重复或近似的段落，从而判断相似度来源是否属于正常模板内容、实验步骤共性内容，还是需要重点关注的异常重复内容。")
    add_screenshot(doc, "teacher01查重结果-相似句子详情.png", "图7-7 相似句子详情界面")
    heading(doc, "7.7 结果导出与演示数据", 2)
    add_p(doc, "系统支持将查重结果导出为 CSV 文件，导出内容包括结果编号、查重任务编号、报告 A 学生、报告 A 文件、报告 B 学生、报告 B 文件、余弦相似度、SimHash 相似度、综合相似度、风险等级和创建时间。系统还提供多班级、多学生、多任务样例数据，以及 task1-txt、task2-docx 和 task3-pdf 三组生成样例报告，用于答辩演示。")
    add_screenshot(doc, "teacher01统计分析.png", "图7-8 查重统计分析界面")

    heading(doc, "8 系统测试与结果分析", 1)
    add_p(doc, "本章对系统运行结果进行验证。测试不仅关注项目是否能够启动，还关注登录、上传、解析、查重、相似句查看和结果导出等主流程是否连贯，同时通过单元测试验证算法和文本解析工具的基本正确性。")
    heading(doc, "8.1 测试环境", 2)
    add_p(doc, "测试工作需要先明确运行环境，否则测试结果缺乏复现基础。本系统测试环境为 Windows 操作系统，后端使用 JDK 17、Maven 3.9.x 和 Spring Boot 3.3.7，前端使用 Node.js、Vue 3、Vite、Element Plus 和 ECharts，数据库使用 MySQL 8.0。测试方式包括 Maven Test、Vite Build 和浏览器手工流程测试。这些环境与项目部署说明保持一致，能够覆盖课程设计答辩时的本机运行场景。")
    heading(doc, "8.2 功能测试", 2)
    add_p(doc, "功能测试围绕教师端主流程和管理员、学生的典型操作展开。测试内容包括用户登录、注册、找回密码、角色菜单显示、课程管理、班级学生分配、实验任务管理、txt/docx/pdf 报告上传、学生班级归属校验、批量上传自动匹配学生、查重任务创建、相似句详情查看、删除失败中文提示和 CSV 结果导出。测试结果表明，系统已经能够支撑从基础数据维护到查重结果导出的完整业务闭环，满足课程设计中“可运行、可演示、可复核”的要求。")
    heading(doc, "8.3 算法与工具类测试", 2)
    add_p(doc, "后端新增了 SimilarityCalculatorTest 和 TextExtractUtilTest，用于验证核心算法和文本解析工具。测试覆盖完全相同文本、不相关文本、空文本、风险等级阈值、相似句匹配、兜底相似句候选、UTF-8 文本解析、GBK 文本解析、docx 文本解析、文字型 pdf 解析、空白 pdf 处理和不支持文件类型异常。")
    add_p(doc, "由于查重算法和文本解析直接决定系统结果是否可信，本文将这两部分从普通功能测试中单独说明。SimilarityCalculatorTest 覆盖相似度计算、空文本处理、风险等级和相似句匹配；TextExtractUtilTest 覆盖 UTF-8、GBK、docx、pdf 和空白 pdf 等场景；ReportCheckApplicationTests 用于验证 Spring Boot 上下文能够正常启动。自动化测试虽然不能覆盖所有真实报告场景，但能够验证核心计算逻辑的基本正确性，尤其是空文本处理、风险阈值和编码兼容等容易被忽略的边界问题。")
    heading(doc, "8.4 构建验证", 2)
    add_p(doc, "后端执行 mvn test 后，测试数随新增 docx/pdf 解析和相似句兜底测试增加，Failures 为 0，Errors 为 0，构建成功。前端执行 npm run build 后，Vite 构建成功，生成 dist 目录。构建过程中存在 chunk 较大的提示，主要由 Element Plus 和 ECharts 等依赖集中打包造成，不影响系统运行和答辩演示。")
    add_p(doc, "系统主流程经验证可以完成登录、查看任务、上传报告、文本解析、发起查重、查看结果、查看相似句、统计分析和导出 CSV 等操作，满足课程设计终稿要求。")

    heading(doc, "9 总结与展望", 1)
    add_p(doc, "本章对课程设计成果进行总结，并说明当前系统仍可继续完善的方向。总结部分重点说明系统已经完成的真实功能和开题报告目标的对应关系，展望部分则如实说明算法、安全性、前端结构和报告导出等方面的后续优化空间。")
    heading(doc, "9.1 总结", 2)
    add_p(doc, "本文设计并实现了一套面向高校实验教学场景的实验报告智能查重与分析系统。系统采用 Spring Boot、Vue 3 和 MySQL 技术栈，围绕课程、班级、实验任务、报告文件和查重结果建立业务模型，完成了管理员、教师和学生三类角色的主要功能。系统支持 txt、docx 和 pdf 报告上传与文本解析，支持教师批量上传报告并自动匹配学生，能够对同一实验任务下的报告进行两两相似度计算，并展示综合相似度、风险等级和相似句子。")
    add_p(doc, "在算法方面，系统采用基于 TF-IDF 思想的文本向量化方法，结合余弦相似度和 SimHash 文本指纹计算综合相似度。系统以同一实验任务下所有已解析报告作为语料集合计算 IDF，使算法流程与开题报告中的设计目标保持一致。测试结果表明，系统能够支撑完整演示流程，具有较好的实用性和可解释性。")
    heading(doc, "9.2 不足与展望", 2)
    add_p(doc, "受课程设计周期限制，系统仍存在一些可以继续完善的地方。第一，当前中文处理采用轻量级二元词条切分，不具备成熟中文分词和深度语义理解能力，后续可以接入 Jieba、HanLP 或基于向量模型的语义相似度算法。第二，当前密码为基础明文存储方式，后续可以引入 BCrypt 提升认证安全性。第三，查重报告当前支持 CSV 导出，后续可以增加 Word 或 PDF 详情报告导出。第四，前端页面主要集中在一个 DashboardView 中，后续可以拆分为独立路由页面，提高代码可维护性。第五，可以继续增加操作日志、历史报告库查重和跨任务查重等功能。")
    add_p(doc, "总体而言，本系统已经完成开题报告中提出的核心目标，能够作为课程设计成果进行演示和答辩。系统定位为辅助教师复核实验报告的轻量级工具，后续可在算法精度、安全性、工程结构和产品化体验方面继续优化。")

    doc.add_page_break()
    center_heading1(doc, "致    谢")
    add_p(doc, "在本次计算机综合课程设计过程中，我围绕实验报告查重这一教学场景完成了需求分析、系统设计、数据库建模、前后端开发、查重算法实现、测试验证和文档整理等工作。通过该项目，我进一步理解了 Web 系统开发流程，也加深了对文本相似度算法和工程实践的认识。")
    add_p(doc, "感谢指导教师在课程设计选题、技术路线和报告撰写方面给予的指导，也感谢课程学习过程中提供帮助的同学和相关开源技术社区。正是这些支持使我能够较为完整地完成本系统设计与实现。由于个人能力和时间有限，系统仍存在不足，后续将继续从算法效果、系统安全性和用户体验等方面进行完善。")

    doc.add_page_break()
    center_heading1(doc, "参考文献")
    refs = [
        "[1] Salton G, Buckley C. Term-weighting approaches in automatic text retrieval[J]. Information Processing & Management, 1988, 24(5): 513-523.",
        "[2] Manku G S, Jain A, Das Sarma A. Detecting near-duplicates for web crawling[C]. Proceedings of the 16th International Conference on World Wide Web, 2007: 141-150.",
        "[3] 李智慧. Spring Boot 企业级开发教程[M]. 北京: 人民邮电出版社, 2021.",
        "[4] 尤雨溪. Vue.js 设计与实现[M]. 北京: 人民邮电出版社, 2022.",
        "[5] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "[6] 宗成庆. 统计自然语言处理[M]. 北京: 清华大学出版社, 2013.",
        "[7] Halim J, Lasut D. Document Plagiarism Detection Application Using Web-Based TF-IDF and Cosine Similarity Methods[J]. bit-Tech, 2024, 7(2): 202-213.",
        "[8] Li Z, Zhu J, Cheng X, Lu Q. Optimized Algorithm Design for Text Similarity Detection Based on Artificial Intelligence and Natural Language Processing[J]. Procedia Computer Science, 2023, 228: 195-202.",
        "[9] Yuan L, Gao S, Pan P. CTSARF: A Chinese Text Similarity Analysis Model based on Residual Fusion[J]. Neurocomputing, 2023, 559: 126801.",
        "[10] Saeed S, Rajput Q, Haider S. SUMEX: A hybrid framework for Semantic textUal siMilarity and EXplanation generation[J]. Information Processing & Management, 2024, 61(5): 103771.",
        "[11] Apache Software Foundation. Apache POI Documentation[EB/OL]. https://poi.apache.org/.",
        "[12] MyBatis-Plus Team. MyBatis-Plus Documentation[EB/OL]. https://baomidou.com/.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        set_run_font(r, 10.5)
        set_para(p, line=1.2, after=3, first=False)

    doc.add_page_break()
    center_heading1(doc, "附    件")
    heading(doc, "附件 A 系统运行与演示说明", 2)
    add_p(doc, "系统源码提交目录包含 backend、frontend、database 和 samples 四个主要部分。后端启动命令为进入 backend 目录后执行 mvn spring-boot:run；前端启动命令为进入 frontend 目录后先执行 npm install，再执行 npm run dev。数据库初始化使用 database/report_check_system_full.sql，该脚本已经包含建库、建表、基础账号和演示数据。")
    heading(doc, "附件 B 样例报告与测试数据", 2)
    add_p(doc, "样例报告位于 samples/generated 目录，其中 task1-txt、task2-docx 和 task3-pdf 可用于教师端批量上传演示。统一数据库脚本 database/report_check_system_full.sql 包含 3 名教师、3 个班级、12 名学生、3 门课程、3 个主演示实验任务以及预置查重结果和相似句详情。")
    heading(doc, "附件 C 主要源码文件说明", 2)
    add_p(doc, "系统主要源码文件包括 SimilarityCalculator.java、CheckTaskService.java、ReportFileService.java、TextExtractUtil.java、DashboardView.vue、LoginView_提交增强版.vue 和 database/report_check_system_full.sql。它们分别对应相似度计算、查重任务生成、报告上传解析、文本提取、前端后台主界面、登录注册找回密码界面以及数据库建表和初始化演示数据。")
    add_p(doc, "后端代码按照 controller、service、mapper、entity、dto、vo、algorithm 和 util 等包组织，前端代码按照接口封装、路由、视图和样式组织。这样的目录结构能够让业务接口、算法实现、文件解析和前端交互保持相对独立，也便于在课程设计答辩时按模块说明代码实现。")

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "计算机综合课程设计报告"
    doc.core_properties.author = "（请填写）"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
