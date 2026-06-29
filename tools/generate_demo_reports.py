from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "samples" / "generated"
SQL_OUT = ROOT / "database" / "demo_seed.sql"

STUDENTS = [
    (101, "20260001", "李明"),
    (102, "20260002", "王芳"),
    (103, "20260003", "赵强"),
    (104, "20260004", "陈晨"),
    (105, "20260005", "孙悦"),
    (106, "20260006", "周杰"),
    (107, "20260007", "吴敏"),
    (108, "20260008", "郑磊"),
    (109, "20260009", "许阳"),
    (110, "20260010", "唐宁"),
    (111, "20260011", "何欣"),
    (112, "20260012", "罗杰"),
]

TASKS = [
    {
        "id": 1,
        "dir": "task1-txt",
        "ext": "txt",
        "title": "实验一-Web系统设计",
        "topic": "Web系统设计与实现",
        "format": "txt",
    },
    {
        "id": 2,
        "dir": "task2-docx",
        "ext": "docx",
        "title": "实验二-文本相似度算法",
        "topic": "文本相似度算法分析",
        "format": "docx",
    },
    {
        "id": 3,
        "dir": "task3-pdf",
        "ext": "pdf",
        "title": "实验三-软件测试与质量分析",
        "topic": "软件测试与质量分析",
        "format": "pdf",
    },
]

COMMON_SENTENCES = [
    "系统采用前后端分离结构，前端负责交互展示，后端负责业务校验、文本解析和查重计算。",
    "报告上传后需要先保存原始文件，再提取正文内容，并记录解析状态、字数和句子数量。",
    "查重过程会对同一实验任务下解析成功的报告进行两两比较，计算综合相似度并保存风险等级。",
    "相似句详情能够帮助教师定位重复段落，比只给出一个总分更容易解释查重结论。",
    "当删除课程、班级、任务或报告时，系统应先检查关联数据，并给出明确中文原因提示。",
]

TASK_PARAGRAPHS = {
    1: [
        "本实验围绕实验报告智能查重系统完成 Web 应用设计，主要包括登录注册、实验任务管理、报告上传、查重结果展示和统计分析等模块。",
        "数据库设计以用户、角色、课程、班级、实验任务、报告文件、查重任务、查重结果和相似句子为核心实体。",
        "前端使用 Vue 和 Element Plus 构建页面，后端使用 Spring Boot 提供 REST 接口，并通过 MyBatis-Plus 访问 MySQL 数据库。",
    ],
    2: [
        "本实验重点分析文本相似度算法，系统使用分词后的词频向量计算余弦相似度，同时使用 SimHash 表示文本指纹。",
        "综合相似度按照余弦相似度和 SimHash 相似度加权得到，用于区分高风险、中风险、低风险和正常报告。",
        "句子级相似度用于生成相似句详情，系统会保存若干组最接近的自然句，便于教师复核。",
    ],
    3: [
        "本实验从软件测试角度验证查重系统，测试范围覆盖账号认证、文件上传、文本解析、查重计算、结果详情和导出功能。",
        "测试用例包含正常流程、异常输入、权限边界、外键删除失败提示以及不同风险等级的报告对比。",
        "测试结果表明，系统能够根据角色控制功能入口，并对解析成功的报告生成相似度结果和相似句明细。",
    ],
}

VARIANT_NOTES = [
    "本报告保留了较多公共句子，适合作为高风险重复案例。",
    "本报告在公共句基础上进行了部分改写，适合作为中风险案例。",
    "本报告只有少量设计思路相近，适合作为低风险案例。",
    "本报告采用独立描述角度，适合作为正常案例。",
]


def paragraph_for(task, index):
    task_id = task["id"]
    student_no, name = STUDENTS[index][1], STUDENTS[index][2]
    topic = task["topic"]
    intro = f"{topic}实验报告由{name}（{student_no}）完成，报告围绕课程设计系统的真实功能进行分析。"
    body = TASK_PARAGRAPHS[task_id]
    if index in (0, 1):
        shared = COMMON_SENTENCES
        note = VARIANT_NOTES[0]
    elif index in (2, 3):
        shared = [
            COMMON_SENTENCES[0],
            "报告提交之后，系统需要保存文件、抽取文本，并把解析状态、字数与句子数量写入数据库。",
            COMMON_SENTENCES[3],
        ]
        note = VARIANT_NOTES[1]
    elif index in (4, 5):
        shared = [
            COMMON_SENTENCES[3],
            "系统在删除已有业务关联的数据时，需要返回清晰的中文提示，避免用户看到底层数据库错误。",
        ]
        note = VARIANT_NOTES[2]
    else:
        shared = [
            f"本报告从{name}的实现过程出发，重点说明{topic}中的功能分工、数据流转和测试结果。",
            "系统演示时需要准备充足数据，确保教师能够看到报告列表、风险等级、统计图和导出结果。",
        ]
        note = VARIANT_NOTES[3]
    conclusion = [
        "通过本次实验，我理解了课程设计系统从需求分析到功能验证的完整流程。",
        "后续可以继续完善密码加密、细粒度权限、批量上传结果反馈和扫描 PDF 的 OCR 处理能力。",
        "综合来看，系统已经能够支撑课堂实验报告提交、查重分析和答辩演示的基本需求。",
    ]
    paragraphs = [intro, *body, *shared, note, *conclusion]
    return "\n".join(paragraphs)


def write_txt(path, text):
    path.write_text(text, encoding="utf-8")


def write_docx(path, text):
    doc = Document()
    doc.add_heading(path.stem, level=1)
    for paragraph in text.split("\n"):
        doc.add_paragraph(paragraph)
    doc.save(path)


def write_pdf(path, text):
    font_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    font_name = "SimHei"
    if font_path.exists():
        pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
    else:
        font_name = "Helvetica"
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    y = height - 54
    c.setFont(font_name, 15)
    c.drawString(54, y, path.stem)
    y -= 34
    c.setFont(font_name, 11)
    for paragraph in text.split("\n"):
        line = ""
        for char in paragraph:
            candidate = line + char
            if c.stringWidth(candidate, font_name, 11) > width - 108:
                c.drawString(54, y, line)
                y -= 20
                line = char
                if y < 60:
                    c.showPage()
                    c.setFont(font_name, 11)
                    y = height - 54
            else:
                line = candidate
        if line:
            c.drawString(54, y, line)
            y -= 20
        y -= 8
        if y < 60:
            c.showPage()
            c.setFont(font_name, 11)
            y = height - 54
    c.save()


def sql_text(value):
    return value.replace("\\", "\\\\").replace("'", "''")


def generate_reports():
    records = []
    for task in TASKS:
        task_dir = OUT / task["dir"]
        task_dir.mkdir(parents=True, exist_ok=True)
        for index, (student_id, student_no, name) in enumerate(STUDENTS):
            text = paragraph_for(task, index)
            filename = f"{student_no}-{name}-{task['title']}.{task['ext']}"
            path = task_dir / filename
            if task["ext"] == "txt":
                write_txt(path, text)
            elif task["ext"] == "docx":
                write_docx(path, text)
            else:
                write_pdf(path, text)
            records.append((task, index, student_id, path, text))
    return records


def generate_demo_seed(records):
    lines = [
        "SET NAMES utf8mb4;",
        "USE report_check_system;",
        "SET FOREIGN_KEY_CHECKS = 0;",
        "DELETE FROM similar_sentence;",
        "DELETE FROM check_result;",
        "DELETE FROM check_task;",
        "DELETE FROM report_file;",
        "SET FOREIGN_KEY_CHECKS = 1;",
        "",
        "INSERT INTO report_file (id, task_id, student_id, original_name, stored_name, file_path, file_type, file_size, parsed_text, word_count, parse_status, parse_message, upload_time) VALUES",
    ]
    report_values = []
    report_id_map = {}
    for task, index, student_id, path, text in records:
        rid = 10000 + task["id"] * 100 + index + 1
        report_id_map[(task["id"], index)] = rid
        rel_path = path.relative_to(ROOT).as_posix()
        report_values.append(
            f"({rid}, {task['id']}, {student_id}, '{sql_text(path.name)}', '{sql_text(path.name)}', "
            f"'{sql_text(rel_path)}', '{task['ext']}', {path.stat().st_size}, '{sql_text(text)}', "
            f"{len(text.replace(chr(10), ''))}, 1, '解析成功，演示数据已预置相似自然句', '2026-05-30 10:00:00')"
        )
    lines.append(",\n".join(report_values) + ";")
    lines.extend([
        "",
        "INSERT INTO check_task (id, experiment_task_id, algorithm, report_count, status, start_time, end_time, created_by, created_time) VALUES",
        "(9001, 1, 'TFIDF_COSINE_SIMHASH', 12, 2, '2026-05-30 10:10:00', '2026-05-30 10:10:08', 2, '2026-05-30 10:10:00'),",
        "(9002, 2, 'TFIDF_COSINE_SIMHASH', 12, 2, '2026-05-30 10:20:00', '2026-05-30 10:20:08', 3, '2026-05-30 10:20:00'),",
        "(9003, 3, 'TFIDF_COSINE_SIMHASH', 12, 2, '2026-05-30 10:30:00', '2026-05-30 10:30:08', 4, '2026-05-30 10:30:00');",
        "",
        "INSERT INTO check_result (id, check_task_id, source_report_id, target_report_id, cosine_similarity, simhash_similarity, final_similarity, risk_level, created_time) VALUES",
    ])
    result_values = []
    result_pairs = []
    scores = [
        ((0, 1), "0.9100", "0.8900", "0.9040", "高风险"),
        ((2, 3), "0.7200", "0.6800", "0.7080", "中风险"),
        ((4, 5), "0.4800", "0.4300", "0.4650", "低风险"),
        ((6, 7), "0.1800", "0.2200", "0.1920", "正常"),
    ]
    result_id = 9100
    for task in TASKS:
        for pair, cosine, simhash, final, risk in scores:
            result_id += 1
            result_pairs.append((result_id, task["id"], pair, risk))
            result_values.append(
                f"({result_id}, {9000 + task['id']}, {report_id_map[(task['id'], pair[0])]}, "
                f"{report_id_map[(task['id'], pair[1])]}, {cosine}, {simhash}, {final}, '{risk}', '2026-05-30 10:40:00')"
            )
    lines.append(",\n".join(result_values) + ";")
    lines.append("")
    lines.append("INSERT INTO similar_sentence (check_result_id, source_sentence, target_sentence, sentence_similarity, created_time) VALUES")
    sentence_values = []
    for rid, task_id, pair, risk in result_pairs:
        if risk == "正常":
            continue
        first = 5 if risk == "高风险" else 3 if risk == "中风险" else 2
        for sentence in COMMON_SENTENCES[:first]:
            sentence_values.append(
                f"({rid}, '{sql_text(sentence)}', '{sql_text(sentence)}', "
                f"{'0.9200' if risk == '高风险' else '0.7100' if risk == '中风险' else '0.5200'}, '2026-05-30 10:40:00')"
            )
    lines.append(",\n".join(sentence_values) + ";")
    SQL_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    records = generate_reports()
    generate_demo_seed(records)
    print(f"generated {len(records)} reports")
    print(f"wrote {SQL_OUT}")
